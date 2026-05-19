"""
Insert 4 blog admin screenshots into docx ABOVE captions Hình 3.32 - 3.35.
"""
import os, sys, io, shutil
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

DOC = r'C:\Users\Kieu Anh\Desktop\CD1\Mã đề 18_Nhóm 3_Bài Thi.docx'
SCREENSHOTS = r'C:\Users\Kieu Anh\.claude\projects\C--xampp-htdocs-huonghoaxinh\screenshots'
BACKUP = r'C:\Users\Kieu Anh\Desktop\CD1\Mã đề 18_Nhóm 3_Bài Thi.beforeBlogImages.docx'

# Map: caption text prefix → screenshot file
INSERTIONS = [
    ("Hình 3.32.", "hinh_3_32_blog_posts_list.png"),
    ("Hình 3.33.", "hinh_3_33_blog_post_form.png"),
    ("Hình 3.34.", "hinh_3_34_blog_categories.png"),
    ("Hình 3.35.", "hinh_3_35_blog_category_form.png"),
]


def insert_image_before(target_para, image_path, width_inches=6.0):
    """Insert a new paragraph (with image inside) BEFORE target_para."""
    # Create empty new <w:p> element BEFORE target
    new_p = OxmlElement('w:p')
    target_para._p.addprevious(new_p)
    new_para = Paragraph(new_p, target_para._parent)
    new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = new_para.add_run()
    run.add_picture(image_path, width=Inches(width_inches))


def main():
    shutil.copy2(DOC, BACKUP)
    print(f"✓ Backup → {BACKUP}\n")

    doc = Document(DOC)

    inserted = 0
    for prefix, screenshot_name in INSERTIONS:
        screenshot_path = os.path.join(SCREENSHOTS, screenshot_name)
        if not os.path.exists(screenshot_path):
            print(f"  ✗ MISSING: {screenshot_path}")
            continue

        # Find caption paragraph (last occurrence — in body, not Danh mục)
        target = None
        for p in doc.paragraphs:
            if p.text.strip().startswith(prefix):
                target = p   # keep updating to get the last match
        if target is None:
            print(f"  ✗ {prefix} caption not found")
            continue

        insert_image_before(target, screenshot_path, width_inches=6.0)
        inserted += 1
        print(f"  ✓ Inserted {screenshot_name} above '{target.text[:60]}'")

    doc.save(DOC)
    print(f"\n✓ Inserted {inserted} blog admin screenshots into docx")


if __name__ == "__main__":
    main()
