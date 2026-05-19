"""
Update Mã đề 18_Nhóm 3_Bài Thi.docx:
  1. Bổ sung 17 entries còn thiếu trong Danh mục hình ảnh (Hình 2.29 - 2.45)
  2. Cập nhật section 3.1.3 — cây cấu trúc thư mục: thêm các file mới
     (BlogPostController, BlogCategoryController, ContactMessageController frontend,
      admin/blog_posts, admin/blog_categories, frontend/account/contact, migrations mới)
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

DOC = r"C:\Users\Kieu Anh\Desktop\CD1\Mã đề 18_Nhóm 3_Bài Thi.docx"


# ============================================================
# 1. Bổ sung 17 entries vào Danh mục hình ảnh
# ============================================================
MISSING_ENTRIES = [
    "Hình 2.29. Wireframe Admin – Danh mục sản phẩm",
    "Hình 2.30. Wireframe Admin – Thêm danh mục",
    "Hình 2.31. Wireframe Admin – Sửa danh mục",
    "Hình 2.32. Wireframe Admin – Khách hàng",
    "Hình 2.33. Wireframe Admin – Form khách hàng",
    "Hình 2.34. Wireframe Admin – Chi tiết khách hàng",
    "Hình 2.35. Wireframe Admin – Đơn hàng",
    "Hình 2.36. Wireframe Admin – Chi tiết đơn hàng",
    "Hình 2.37. Wireframe Admin – Voucher",
    "Hình 2.38. Wireframe Admin – Thêm voucher",
    "Hình 2.39. Wireframe Admin – Sửa voucher",
    "Hình 2.40. Wireframe Admin – Đánh giá",
    "Hình 2.41. Wireframe Admin – Chat với khách hàng (chat inbox)",
    "Hình 2.42. Wireframe Admin – Chi tiết tin nhắn",
    "Hình 2.43. Wireframe Admin – Thống kê doanh thu",
    "Hình 2.44. Wireframe Admin – Cài đặt website",
    "Hình 2.45. Wireframe Admin – Hồ sơ admin",
]


def update_danh_muc(doc):
    """Insert Hình 2.29 - 2.45 entries after Hình 2.28 in Danh mục hình ảnh."""
    target = None
    for i, p in enumerate(doc.paragraphs):
        if i > 130: break
        if p.text.strip().startswith("Hình 2.28."):
            target = p
            break
    if target is None:
        print("  KHÔNG TÌM THẤY Hình 2.28 trong Danh mục"); return

    # Insert in reverse to preserve order
    for entry in reversed(MISSING_ENTRIES):
        new_p = OxmlElement('w:p')
        target._p.addnext(new_p)
        para = Paragraph(new_p, target._parent)
        run = para.add_run(entry)
        run.font.name = "Times New Roman"
        run.font.size = Pt(13)
        para.paragraph_format.line_spacing = 1.5
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    print(f"  ✓ Đã thêm {len(MISSING_ENTRIES)} entries vào Danh mục hình ảnh (Hình 2.29 → 2.45)")


# ============================================================
# 2. Cập nhật cây thư mục ở section 3.1.3 (Table index = 6)
# ============================================================
NEW_TREE = """huonghoaxinh/
├── app/
│   ├── Http/
│   │   ├── Controllers/
│   │   │   ├── Admin/                       # Controllers khu vực quản trị
│   │   │   │   ├── DashboardController.php
│   │   │   │   ├── ProductController.php
│   │   │   │   ├── CategoryController.php
│   │   │   │   ├── OrderController.php
│   │   │   │   ├── UserController.php
│   │   │   │   ├── VoucherController.php
│   │   │   │   ├── ReviewController.php
│   │   │   │   ├── RevenueController.php
│   │   │   │   ├── ContactMessageController.php
│   │   │   │   ├── BlogPostController.php       ★ MỚI - CRUD bài viết blog (Quill)
│   │   │   │   ├── BlogCategoryController.php   ★ MỚI - CRUD danh mục blog
│   │   │   │   ├── WebsiteSettingController.php
│   │   │   │   ├── AdminProfileController.php
│   │   │   │   └── Auth/
│   │   │   ├── Frontend/                    # Controllers khu vực khách hàng
│   │   │   │   ├── HomeController.php
│   │   │   │   ├── ShopController.php
│   │   │   │   ├── ProductController.php
│   │   │   │   ├── CartController.php
│   │   │   │   ├── PaymentController.php
│   │   │   │   ├── WishlistController.php
│   │   │   │   ├── ProductReviewController.php
│   │   │   │   ├── BlogController.php
│   │   │   │   ├── ContactMessageController.php ★ MỚI - Chat với cửa hàng
│   │   │   │   └── ShippingEstimateController.php
│   │   │   └── Auth/                        # Controllers xác thực (Breeze)
│   │   └── Middleware/                      # Middleware phân quyền (admin / auth)
│   ├── Models/                              # Eloquent Model ánh xạ bảng DB
│   │   ├── User.php, Admin.php
│   │   ├── Product.php, Category.php
│   │   ├── Order.php, OrderItem.php           ★ Order có thêm: sender_*, recipient_*,
│   │   │                                        delivery_date, delivery_time_slot,
│   │   │                                        recipient_message
│   │   ├── Cart.php, Wishlist.php
│   │   ├── Voucher.php, VoucherUserUsage.php
│   │   ├── Review.php
│   │   ├── BlogPost.php, BlogCategory.php
│   │   ├── ContactMessage.php                ★ Có thêm user_id (1 user = 1 chat thread)
│   │   ├── ContactReply.php                  ★ Có thêm user_id, customer_read_at
│   │   └── WebsiteSetting.php
│   ├── Services/                            # Lớp xử lý nghiệp vụ tách khỏi controller
│   └── Providers/                           # Service Providers
├── resources/
│   ├── views/
│   │   ├── admin/                           # Blade views khu vực quản trị
│   │   │   ├── layouts/
│   │   │   ├── dashboard.blade.php
│   │   │   ├── products/, categories/
│   │   │   ├── orders/                      # Order detail hiển thị sender + recipient
│   │   │   ├── users/, vouchers/, reviews/
│   │   │   ├── blog_posts/                   ★ MỚI: index + form (Quill rich editor)
│   │   │   ├── blog_categories/              ★ MỚI: index + form CRUD
│   │   │   ├── revenue/, settings/
│   │   │   ├── contact_messages/            # Inbox chat + thread detail (realtime)
│   │   │   └── profile/
│   │   ├── frontend/                        # Blade views khu vực khách hàng
│   │   │   ├── layouts/                     # Đã thêm floating chat FAB + badge
│   │   │   ├── home.blade.php, shop.blade.php
│   │   │   ├── about.blade.php, contact.blade.php
│   │   │   ├── product/, cart/
│   │   │   ├── checkout/                    ★ Sender + Recipient + thiệp + giao
│   │   │   ├── account/                     # Hồ sơ
│   │   │   │   └── contact/                  ★ MỚI: index + show (chat realtime)
│   │   │   ├── orders/
│   │   │   ├── blog/                        # Magazine layout + reading progress
│   │   │   ├── wishlist/, reviews/
│   │   │   └── vouchers.blade.php
│   │   └── auth/                            # Form đăng nhập / đăng ký / quên mật khẩu
│   ├── css/                                 # SCSS / CSS nguồn (bundle bằng Vite)
│   └── js/                                  # JavaScript nguồn
├── routes/
│   ├── web.php                              # Toàn bộ HTTP routes (kể cả /admin/blog-*)
│   └── auth.php                             # Routes của Laravel Breeze
├── database/
│   ├── migrations/                          # Định nghĩa schema (CREATE TABLE…)
│   │   ├── ...                              # Các migration gốc
│   │   ├── 2026_05_18_000000_add_user_id_to_contact_replies.php       ★
│   │   ├── 2026_05_18_010000_add_user_id_to_contact_messages.php       ★
│   │   ├── 2026_05_18_020000_add_customer_read_at_to_contact_replies.php ★
│   │   └── 2026_05_18_030000_add_recipient_fields_to_orders.php       ★
│   ├── seeders/                             # Dữ liệu mẫu (admin, sản phẩm, blog…)
│   └── factories/                           # Sinh dữ liệu test
├── public/
│   ├── index.php                            # Entry-point web
│   ├── img/                                 # Ảnh tĩnh đi kèm dự án
│   └── storage/                             # Symlink → storage/app/public
├── storage/
│   ├── app/public/                          # File upload (ảnh sản phẩm, logo, ảnh blog…)
│   ├── framework/                           # Cache, views compiled, session
│   └── logs/                                # Log Laravel
├── config/                                  # Cấu hình toàn cục (app, db, mail…)
├── tests/                                   # Unit / Feature test PHPUnit
├── .env                                     # Biến môi trường (DB, mail, APP_KEY…)
├── composer.json                            # Khai báo PHP dependencies
├── package.json                             # Khai báo NPM dependencies
├── vite.config.js                           # Cấu hình bundler Vite
└── artisan                                  # Laravel CLI (migrate, serve, tinker…)"""


def update_tree(doc):
    """Replace content of the tree code-block table with NEW_TREE."""
    target_table = None
    for i, t in enumerate(doc.tables):
        if not t.rows: continue
        if len(t.rows) == 1 and len(t.columns) == 1:
            cell_text = t.rows[0].cells[0].text
            if cell_text.startswith("huonghoaxinh/"):
                target_table = t
                break
    if target_table is None:
        print("  KHÔNG TÌM THẤY table tree"); return

    cell = target_table.rows[0].cells[0]

    # Wipe existing paragraphs in cell
    tc = cell._tc
    for p in list(tc):
        if p.tag.endswith('}p'):
            tc.remove(p)

    # Add new paragraphs with new tree
    lines = NEW_TREE.split('\n')
    for line in lines:
        new_p = OxmlElement('w:p')
        tc.append(new_p)
        para = Paragraph(new_p, cell)
        run = para.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(10)
        para.paragraph_format.line_spacing = 1.15
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    print(f"  ✓ Đã cập nhật cây cấu trúc thư mục ({len(lines)} dòng) với các file mới")


# ============================================================
# 3. Bổ sung paragraph giải thích các file mới vào section 3.1.3
# ============================================================
NEW_INTRO_NOTE = (
    "So với phiên bản ban đầu, dự án đã bổ sung các file mới (đánh dấu ★) "
    "cho 3 nhóm chức năng: (1) Blog CMS phía admin gồm BlogPostController, "
    "BlogCategoryController cùng các view tương ứng với trình soạn thảo Quill; "
    "(2) Chat 2 chiều khách – admin với Frontend\\ContactMessageController và "
    "view account/contact lưu trực tiếp trong DB (không qua email); (3) Thông "
    "tin người gửi / người nhận trên đơn hàng (sender_*, recipient_*, "
    "delivery_date, recipient_message) cùng 4 migration ngày 2026_05_18."
)


def add_intro_note(doc):
    """Insert explanatory paragraph after the first paragraph of section 3.1.3."""
    h313_idx = None
    paras = list(doc.paragraphs)
    for i, p in enumerate(paras):
        if p.style.name == "Heading 3" and p.text.strip().startswith("3.1.3"):
            h313_idx = i; break
    if h313_idx is None:
        print("  KHÔNG TÌM THẤY section 3.1.3"); return

    # Find first non-empty para after heading
    intro_para = None
    intro_idx = None
    for i in range(h313_idx + 1, h313_idx + 6):
        if i < len(paras) and paras[i].text.strip():
            intro_para = paras[i]
            intro_idx = i; break
    if intro_para is None:
        print("  KHÔNG TÌM THẤY intro para của 3.1.3"); return

    # Check if note already exists to avoid duplicate
    if intro_idx + 1 < len(paras):
        next_text = paras[intro_idx + 1].text.strip()
        if "đánh dấu ★" in next_text:
            print("  Note đã tồn tại — bỏ qua"); return

    # Insert new paragraph AFTER intro
    new_p = OxmlElement('w:p')
    intro_para._p.addnext(new_p)
    para = Paragraph(new_p, intro_para._parent)
    run = para.add_run(NEW_INTRO_NOTE)
    run.font.name = "Times New Roman"
    run.font.size = Pt(13)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    para.paragraph_format.first_line_indent = Pt(18)
    print("  ✓ Đã thêm đoạn giải thích các file mới (đánh dấu ★)")


def main():
    doc = Document(DOC)
    print("=== UPDATE REPORT ===\n")
    print("[1/3] Danh mục hình ảnh:")
    update_danh_muc(doc)
    print("\n[2/3] Cây cấu trúc thư mục 3.1.3:")
    update_tree(doc)
    print("\n[3/3] Đoạn giải thích file mới:")
    add_intro_note(doc)
    doc.save(DOC)
    print("\nOK — đã lưu lại file.")


if __name__ == "__main__":
    main()
