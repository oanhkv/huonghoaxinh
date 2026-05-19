"""
Rewrite v2 sections 3.1.1 and 3.1.2 — concise, report tone, no personal
pronouns, focused on key points.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

DOC = r"C:\Users\Kieu Anh\Desktop\CD1\Mã đề 18_Nhóm 3_Bài Thi.docx"


def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def set_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), '808080')
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def new_para_before(anchor, segments, *,
                    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                    line_spacing=1.5, first_line_indent=None,
                    font_name="Times New Roman", size_pt=13):
    new_p = OxmlElement('w:p')
    anchor._p.addprevious(new_p)
    p = Paragraph(new_p, anchor._parent)
    if isinstance(segments, str):
        segments = [(segments, False, False)]
    for seg in segments:
        if len(seg) == 2:
            txt, bold = seg; italic = False
        else:
            txt, bold, italic = seg
        run = p.add_run(txt)
        run.font.name = font_name
        run.font.size = Pt(size_pt)
        run.bold = bold
        run.italic = italic
    p.alignment = alignment
    pf = p.paragraph_format
    pf.line_spacing = line_spacing
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    return p


def new_bullet_before(anchor, segments):
    if isinstance(segments, str):
        segments = [(segments, False, False)]
    segs = [("•  ", False, False)] + list(segments)
    p = new_para_before(anchor, segs)
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.first_line_indent = Cm(-0.6)
    return p


def new_code_block_before(anchor, code_text, doc):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    anchor._p.addprevious(tbl._tbl)
    cell = tbl.rows[0].cells[0]
    cell.text = ""
    set_cell_shading(cell, "F5F5F5")
    set_cell_borders(cell)
    lines = code_text.rstrip("\n").split("\n")
    cell.paragraphs[0].text = lines[0]
    for line in lines[1:]:
        cell.add_paragraph(line)
    for cp in cell.paragraphs:
        cp.paragraph_format.line_spacing = 1.15
        cp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        cp.paragraph_format.space_before = Pt(0)
        cp.paragraph_format.space_after = Pt(0)
        cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for r in cp.runs:
            r.font.name = "Consolas"
            r.font.size = Pt(11)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    return tbl


def main():
    doc = Document(DOC)
    paras = list(doc.paragraphs)
    h311 = h312 = h313 = None
    for i, p in enumerate(paras):
        if p.style.name == "Heading 3":
            t = p.text.strip()
            if t.startswith("3.1.1"): h311 = (i, p)
            elif t.startswith("3.1.2"): h312 = (i, p)
            elif t.startswith("3.1.3"): h313 = (i, p)

    # ---- Find table inserted in 3.1.2 (code block) so we remove it too ----
    # We'll remove all paragraphs AND tables between (3.1.1, 3.1.3)
    # Approach: walk body XML, drop every <w:p> / <w:tbl> between
    # heading 3.1.1 and 3.1.3.
    body = doc.element.body
    elems = list(body)

    def text_of(el):
        return ''.join(t.text or '' for t in el.iter(qn('w:t')))

    # locate body indexes of the two headings
    idx_h311 = idx_h313 = None
    for k, el in enumerate(elems):
        if el.tag.endswith('}p'):
            t = text_of(el).strip()
            if t.startswith("3.1.1.") and 'Mô hình MVC' in t:
                idx_h311 = k
            elif t.startswith("3.1.3.") and 'Cấu trúc thư mục' in t:
                idx_h313 = k

    if idx_h311 is None or idx_h313 is None:
        raise SystemExit("Không tìm thấy heading 3.1.1 hoặc 3.1.3 trong body.")

    # Remove everything strictly between them (paragraphs + tables)
    for k in range(idx_h313 - 1, idx_h311, -1):
        body.remove(elems[k])

    # ============================================================
    # Insert NEW concise content for 3.1.1 + 3.1.2
    # Anchor = heading 3.1.3
    # ============================================================
    # Re-fetch h313 after deletions
    h313_p = None
    for p in doc.paragraphs:
        if p.style.name == "Heading 3" and p.text.strip().startswith("3.1.3"):
            h313_p = p; break

    # We need to RE-INSERT the 3.1.2 heading (it was deleted because it lay
    # between 3.1.1 and 3.1.3). Same for 3.1.1 heading? — No, 3.1.1 heading
    # is at idx_h311 which is the boundary (not removed). 3.1.2 heading was
    # in between → was removed. We need to re-create it.

    anchor = h313_p

    # ========== 3.1.1 body ==========
    new_para_before(anchor,
        "MVC (Model – View – Controller) là mẫu thiết kế chia ứng dụng "
        "thành ba lớp tách biệt, mỗi lớp giữ một vai trò riêng. Đây cũng "
        "là cấu trúc mặc định được Laravel khuyến nghị áp dụng.",
        first_line_indent=Pt(18))

    new_bullet_before(anchor, [
        ("Model:", True),
        (" lớp dữ liệu, đại diện cho bảng trong MySQL thông qua "
         "Eloquent ORM. Mỗi Model bao gọn cả truy vấn và các quy tắc "
         "nghiệp vụ liên quan đến dữ liệu đó.", False),
    ])
    new_bullet_before(anchor, [
        ("View:", True),
        (" lớp hiển thị, viết bằng Blade Template (", False),
        (".blade.php", False, True),
        ("). Kết hợp HTML với các cú pháp Blade như ", False),
        ("{{ $bien }}", False, True), (", ", False),
        ("@foreach", False, True), (", ", False),
        ("@extends", False, True),
        (". View chỉ hiển thị dữ liệu, không xử lý nghiệp vụ.", False),
    ])
    new_bullet_before(anchor, [
        ("Controller:", True),
        (" lớp điều phối, nhận HTTP Request, gọi Model lấy dữ liệu rồi "
         "trả về View phù hợp.", False),
    ])

    new_para_before(anchor, [
        ("Luồng xử lý một request: ", True),
        ("Trình duyệt → ", False),
        ("routes/web.php", False, True),
        (" → Controller → Model (Eloquent) → trả dữ liệu cho Controller "
         "→ render Blade View → trả HTML về trình duyệt.", False),
    ], first_line_indent=Pt(18))

    # ========== 3.1.2 heading ==========
    h312_new_p = new_para_before(anchor, "3.1.2. Áp dụng MVC vào dự án Hương Hoa Xinh")
    # Apply Heading 3 style
    h312_new_p.style = doc.styles["Heading 3"]

    # ========== 3.1.2 body — concise ==========
    new_para_before(anchor,
        "Dự án Hương Hoa Xinh tuân thủ cấu trúc MVC chuẩn của Laravel 11, "
        "tổ chức theo bốn nhóm chính:",
        first_line_indent=Pt(18))

    new_bullet_before(anchor, [
        ("Model – ", True), ("app/Models/", True, True),
        (": 16 Model ứng với 16 bảng dữ liệu, gồm các nhóm chính: tài "
         "khoản (", False),
        ("User", False, True), (", ", False), ("Admin", False, True),
        ("); sản phẩm (", False),
        ("Product", False, True), (", ", False), ("Category", False, True),
        ("); đơn hàng (", False),
        ("Order", False, True), (", ", False), ("OrderItem", False, True),
        ("); giỏ – yêu thích (", False),
        ("Cart", False, True), (", ", False), ("Wishlist", False, True),
        ("); khuyến mãi (", False),
        ("Voucher", False, True), (", ", False),
        ("VoucherUserUsage", False, True),
        ("); blog (", False),
        ("BlogPost", False, True), (", ", False),
        ("BlogCategory", False, True),
        ("); và các Model bổ trợ (", False),
        ("Review", False, True), (", ", False),
        ("ContactMessage", False, True), (", ", False),
        ("ContactReply", False, True), (", ", False),
        ("WebsiteSetting", False, True), (").", False),
    ])

    new_bullet_before(anchor, [
        ("Controller – ", True), ("app/Http/Controllers/", True, True),
        (": chia 3 nhánh – ", False),
        ("Admin/", False, True),
        (" (11 controller cho khu vực quản trị), ", False),
        ("Frontend/", False, True),
        (" (9 controller cho phía khách hàng) và ", False),
        ("Auth/", False, True),
        (" (controller xác thực sinh sẵn bởi Laravel Breeze).", False),
    ])

    new_bullet_before(anchor, [
        ("View – ", True), ("resources/views/", True, True),
        (": 4 nhóm Blade Template – ", False),
        ("layouts/", False, True), (" (template cha), ", False),
        ("frontend/", False, True), (" (giao diện khách hàng), ", False),
        ("admin/", False, True), (" (giao diện quản trị) và ", False),
        ("auth/", False, True),
        (" (form đăng nhập / đăng ký).", False),
    ])

    new_bullet_before(anchor, [
        ("Route – ", True), ("routes/web.php", True, True),
        (": khai báo toàn bộ URL của website, gắn mỗi URL với một method "
         "Controller cụ thể, đồng thời áp dụng middleware bảo vệ tương "
         "ứng (", False),
        ("auth", False, True), (" cho khách hàng đã đăng nhập, ", False),
        ("admin", False, True), (" cho khu vực quản trị).", False),
    ])

    new_para_before(anchor, "Ví dụ một đoạn khai báo route trong dự án:",
                    first_line_indent=Pt(18))

    new_code_block_before(anchor,
        "Route::get('/shop', [ShopController::class, 'index'])->name('shop');\n"
        "\n"
        "Route::middleware('admin')->prefix('admin')->name('admin.')->group(function () {\n"
        "    Route::resource('products', ProductController::class);\n"
        "    Route::resource('orders',   OrderController::class);\n"
        "    Route::resource('vouchers', VoucherController::class);\n"
        "});",
        doc=doc)

    doc.save(DOC)
    print("OK")


if __name__ == "__main__":
    main()
