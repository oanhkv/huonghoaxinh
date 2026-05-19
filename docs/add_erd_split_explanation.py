"""
Insert explanation paragraphs in section 2.3 about why ERD was split into 3 sub-diagrams
and how the 3 ERDs link via bridge entities (users, products).
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

DOC = r"C:\Users\Kieu Anh\Desktop\CD1\Mã đề 18_Nhóm 3_Bài Thi.docx"


def new_para_before(anchor, segments, *,
                    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                    line_spacing=1.5,
                    first_line_indent=None,
                    style_name=None,
                    font="Times New Roman", size_pt=13):
    new_p = OxmlElement('w:p')
    anchor._p.addprevious(new_p)
    p = Paragraph(new_p, anchor._parent)
    if style_name:
        p.style = anchor.part.document.styles[style_name]
    if isinstance(segments, str):
        segments = [(segments, False, False)]
    for seg in segments:
        if len(seg) == 2: txt, bold = seg; italic = False
        else: txt, bold, italic = seg
        run = p.add_run(txt)
        run.font.name = font
        run.font.size = Pt(size_pt)
        run.bold = bold
        run.italic = italic
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.line_spacing = line_spacing
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    return p


def new_bullet_before(anchor, segments):
    if isinstance(segments, str):
        segments = [(segments, False, False)]
    segs = [("•  ", False, False)] + list(segments)
    p = new_para_before(anchor, segs)
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.first_line_indent = Cm(-0.6)
    return p


def main():
    doc = Document(DOC)

    # Find heading 2.4 to insert before
    anchor = None
    for p in doc.paragraphs:
        if p.style.name == "Heading 2" and p.text.strip().startswith("2.4"):
            anchor = p; break
    if anchor is None:
        print("Không tìm thấy heading 2.4"); return

    # Idempotency check
    for p in doc.paragraphs:
        if "Do hệ thống có tới 16 bảng" in p.text:
            print("Đã có nội dung — bỏ qua để không tạo trùng"); return

    # ===== Heading 3 phụ =====
    h = new_para_before(anchor, "2.3.1. Lý do chia ERD thành 3 sơ đồ và mối liên hệ")
    h.style = doc.styles["Heading 3"]

    # ===== Paragraph 1: Lý do =====
    new_para_before(anchor,
        "Do hệ thống có tới 16 bảng và mật độ quan hệ rất cao (bảng users là "
        "trung tâm với 7 mối quan hệ trực tiếp, bảng products có 5 mối quan "
        "hệ), việc vẽ tất cả lên cùng một sơ đồ ERD sẽ khiến hình bị rối, "
        "đường nối chồng chéo và khó in vừa khổ giấy A4. Vì vậy, nhóm em "
        "chia ERD thành ba sơ đồ con theo từng nhóm chức năng (domain) — vừa "
        "đảm bảo đầy đủ thông tin, vừa dễ trình bày và đọc hiểu.",
        first_line_indent=Pt(18))

    # ===== Paragraph 2: 3 ERDs =====
    new_para_before(anchor,
        "Cụ thể, ba sơ đồ ERD được phân chia như sau:",
        first_line_indent=Pt(18))

    new_bullet_before(anchor, [
        ("Sơ đồ ERD 1 — Khách hàng và Sản phẩm: ", True),
        ("gồm 6 thực thể chính ", False),
        ("users, categories, products, carts, wishlists, reviews", False, True),
        (". Đây là phần lõi của module bán hàng — mọi luồng nghiệp vụ "
         "khác đều bắt đầu từ users hoặc products.", False),
    ])

    new_bullet_before(anchor, [
        ("Sơ đồ ERD 2 — Đơn hàng và Thanh toán: ", True),
        ("gồm 4 thực thể chính ", False),
        ("orders, order_items, vouchers, voucher_user_usages", False, True),
        (". Sơ đồ này tập trung vào luồng đặt hàng — bao gồm cả các "
         "trường mở rộng cho người gửi / người nhận, ngày giao, lời nhắn "
         "thiệp đã được bổ sung gần đây.", False),
    ])

    new_bullet_before(anchor, [
        ("Sơ đồ ERD 3 — Nội dung và Giao tiếp: ", True),
        ("gồm 6 thực thể ", False),
        ("admins, contact_messages, contact_replies, blog_categories, "
         "blog_posts, website_settings", False, True),
        (". Tập hợp các bảng phục vụ quản trị nội dung blog, chat 2 chiều "
         "khách – admin và cấu hình website. Đây là phần đứng độc lập "
         "tương đối với phần bán hàng.", False),
    ])

    # ===== Paragraph 3: Cách liên kết =====
    new_para_before(anchor, [
        ("Mối liên kết giữa ba sơ đồ ", True),
        ("được thực hiện thông qua hai thực thể \"bridge\" — đó là ", False),
        ("users", False, True),
        (" và ", False),
        ("products", False, True),
        (". Cả hai xuất hiện đầy đủ trong ERD 1 (cùng các thuộc tính chi "
         "tiết); ở ERD 2 và ERD 3 chúng được vẽ lại dưới dạng hộp viền "
         "nét đứt với ghi chú [bridge ← ERD 1] để người đọc biết đó "
         "không phải bảng mới mà chỉ là tham chiếu sang sơ đồ khác. "
         "Cụ thể: ERD 2 dùng users.id (làm khoá ngoại trong orders, "
         "voucher_user_usages) và products.id (làm khoá ngoại trong "
         "order_items); ERD 3 dùng users.id (làm khoá ngoại trong "
         "contact_messages, contact_replies).", False),
    ], first_line_indent=Pt(18))

    # ===== Paragraph 4: Tóm lại =====
    new_para_before(anchor,
        "Cách chia này tuân theo nguyên tắc \"high cohesion, low coupling\" "
        "trong thiết kế CSDL: các bảng có liên hệ nghiệp vụ chặt chẽ thì "
        "được nhóm chung một sơ đồ, còn liên kết giữa các nhóm chỉ thông "
        "qua một vài khoá ngoại rõ ràng (users.id, products.id). Khi cần "
        "tra cứu chi tiết các thuộc tính của bảng users hoặc products, "
        "người đọc luôn quay lại ERD 1 — nên không có thông tin nào bị mất "
        "khi tách sơ đồ.",
        first_line_indent=Pt(18))

    doc.save(DOC)
    print("OK — đã thêm 1 heading phụ + 5 đoạn giải thích vào section 2.3")


if __name__ == "__main__":
    main()
