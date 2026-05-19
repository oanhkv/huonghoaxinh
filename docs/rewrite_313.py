"""
Rewrite section 3.1.3 in the report:
  - Remove the screenshot + its 'Hình 3.1' caption.
  - Insert: intro paragraph + monospace directory tree + description table.
  - Re-number all body 'Hình 3.X' (X>=2) → 'Hình 3.(X-1)'.
  - Remove 'Hình 3.1' entry from 'Danh mục hình ảnh' and re-number the rest.
  - Add the missing 'Hồ sơ admin' entry (Hình 3.30 after renumber).
"""
import re
from copy import deepcopy
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

DOC = r"C:\Users\Kieu Anh\Desktop\CD1\Mã đề 18_Nhóm 3_Bài Thi.docx"

# ----------------------------- content -----------------------------
INTRO = (
    "Dự án Hương Hoa Xinh tuân thủ chuẩn cấu trúc thư mục mặc định của "
    "framework Laravel, đồng thời tách rõ phần Frontend (khách hàng) và "
    "Admin (quản trị) ở cả Controllers lẫn Views để dễ bảo trì. Cây thư "
    "mục chính của dự án được thể hiện như sau:"
)

TREE = """huonghoaxinh/
├── app/
│   ├── Http/
│   │   ├── Controllers/
│   │   │   ├── Admin/                       # Controllers khu vực quản trị
│   │   │   │   ├── DashboardController.php
│   │   │   │   ├── ProductController.php
│   │   │   │   ├── CategoryController.php
│   │   │   │   ├── OrderController.php
│   │   │   │   ├── UserController.php
│   │   │   │   ├── VoucherController.php
│   │   │   │   ├── ReviewController.php
│   │   │   │   ├── RevenueController.php
│   │   │   │   ├── ContactMessageController.php
│   │   │   │   ├── WebsiteSettingController.php
│   │   │   │   ├── AdminProfileController.php
│   │   │   │   └── Auth/
│   │   │   ├── Frontend/                    # Controllers khu vực khách hàng
│   │   │   │   ├── HomeController.php
│   │   │   │   ├── ShopController.php
│   │   │   │   ├── ProductController.php
│   │   │   │   ├── CartController.php
│   │   │   │   ├── PaymentController.php
│   │   │   │   ├── WishlistController.php
│   │   │   │   ├── ProductReviewController.php
│   │   │   │   ├── BlogController.php
│   │   │   │   └── ShippingEstimateController.php
│   │   │   └── Auth/                        # Controllers xác thực (Breeze)
│   │   └── Middleware/                      # Middleware phân quyền (admin / auth)
│   ├── Models/                              # Eloquent Model ánh xạ bảng DB
│   │   ├── User.php, Admin.php
│   │   ├── Product.php, Category.php
│   │   ├── Order.php, OrderItem.php
│   │   ├── Cart.php, Wishlist.php
│   │   ├── Voucher.php, VoucherUserUsage.php
│   │   ├── Review.php
│   │   ├── BlogPost.php, BlogCategory.php
│   │   ├── ContactMessage.php, ContactReply.php
│   │   └── WebsiteSetting.php
│   ├── Mail/                                # Mailable (xác nhận đơn, trả lời liên hệ)
│   ├── Services/                            # Lớp xử lý nghiệp vụ tách khỏi controller
│   └── Providers/                           # Service Providers
├── resources/
│   ├── views/
│   │   ├── admin/                           # Blade views khu vực quản trị
│   │   │   ├── layouts/
│   │   │   ├── dashboard.blade.php
│   │   │   ├── products/, categories/
│   │   │   ├── orders/, users/
│   │   │   ├── vouchers/, reviews/
│   │   │   ├── revenue/, settings/
│   │   │   ├── contact_messages/
│   │   │   └── profile/
│   │   ├── frontend/                        # Blade views khu vực khách hàng
│   │   │   ├── layouts/
│   │   │   ├── home.blade.php, shop.blade.php
│   │   │   ├── about.blade.php, contact.blade.php
│   │   │   ├── product/, cart/, checkout/
│   │   │   ├── account/, orders/
│   │   │   ├── blog/, wishlist/, reviews/
│   │   │   └── vouchers.blade.php
│   │   └── auth/                            # Form đăng nhập / đăng ký / quên mật khẩu
│   ├── css/                                 # SCSS / CSS nguồn (bundle bằng Vite)
│   └── js/                                  # JavaScript nguồn
├── routes/
│   ├── web.php                              # Khai báo toàn bộ HTTP routes
│   └── auth.php                             # Routes của Laravel Breeze
├── database/
│   ├── migrations/                          # Định nghĩa schema (CREATE TABLE…)
│   ├── seeders/                             # Dữ liệu mẫu (admin, sản phẩm…)
│   └── factories/                           # Sinh dữ liệu test
├── public/
│   ├── index.php                            # Entry-point web
│   ├── img/                                 # Ảnh tĩnh đi kèm dự án
│   └── storage/                             # Symlink → storage/app/public
├── storage/
│   ├── app/public/                          # File upload (ảnh sản phẩm, logo…)
│   ├── framework/                           # Cache, views compiled, session
│   └── logs/                                # Log Laravel
├── config/                                  # Cấu hình toàn cục (app, db, mail…)
├── tests/                                   # Unit / Feature test PHPUnit
├── .env                                     # Biến môi trường (DB, mail, APP_KEY…)
├── composer.json                            # Khai báo PHP dependencies
├── package.json                             # Khai báo NPM dependencies
├── vite.config.js                           # Cấu hình bundler Vite
└── artisan                                  # Laravel CLI (migrate, serve, tinker…)
"""

DESC_ROWS = [
    ("app/Http/Controllers/Admin/",
     "Chứa controller cho khu vực quản trị: dashboard, sản phẩm, danh mục, đơn hàng, "
     "khách hàng, voucher, đánh giá, doanh thu, tin nhắn liên hệ, cài đặt website."),
    ("app/Http/Controllers/Frontend/",
     "Controller phía khách hàng: trang chủ, shop, chi tiết sản phẩm, giỏ hàng, "
     "thanh toán, wishlist, blog, đánh giá sản phẩm, ước tính phí ship."),
    ("app/Http/Controllers/Auth/",
     "Controller xác thực do Laravel Breeze sinh: đăng ký, đăng nhập, "
     "quên mật khẩu, xác thực email."),
    ("app/Http/Middleware/",
     "Middleware phân quyền (Admin / Auth / Guest), kiểm tra is_locked, "
     "redirect khi chưa login."),
    ("app/Models/",
     "Eloquent Model ánh xạ 1-1 với các bảng trong CSDL ban_hoa "
     "(User, Admin, Product, Category, Order, OrderItem, Cart, Wishlist, "
     "Voucher, Review, BlogPost, ContactMessage, WebsiteSetting…)."),
    ("app/Mail/",
     "Lớp Mailable định nghĩa email gửi đi (xác nhận đơn hàng, "
     "trả lời tin nhắn liên hệ, đặt lại mật khẩu)."),
    ("app/Services/",
     "Lớp dịch vụ tách nghiệp vụ phức tạp khỏi controller "
     "(ví dụ OrderInventoryService xử lý trừ/hoàn tồn kho)."),
    ("resources/views/admin/",
     "Blade template cho trang quản trị: layout, dashboard, các CRUD form, "
     "trang doanh thu, cài đặt, tin nhắn…"),
    ("resources/views/frontend/",
     "Blade template trang khách hàng: home, shop, product, cart, checkout, "
     "account, blog, wishlist, reviews, vouchers."),
    ("resources/views/auth/",
     "Form đăng nhập / đăng ký / quên mật khẩu / xác thực email."),
    ("routes/web.php",
     "Khai báo toàn bộ route HTTP: route khách hàng, nhóm route /admin (có "
     "middleware admin), route Breeze được include từ auth.php."),
    ("database/migrations/",
     "Mỗi file là một migration định nghĩa cấu trúc bảng (CREATE TABLE) "
     "và các cột bổ sung (add_role_to_users, add_is_locked_to_users, …)."),
    ("database/seeders/",
     "Dữ liệu mẫu khi chạy `php artisan db:seed` "
     "(admin mặc định, sản phẩm, danh mục, blog…)."),
    ("public/",
     "Thư mục web-root. Chứa index.php (entry-point), ảnh tĩnh /img, "
     "và symlink /storage trỏ tới file upload."),
    ("storage/app/public/",
     "Nơi lưu file upload thực sự (ảnh sản phẩm, logo, hero, "
     "đính kèm reply liên hệ)."),
    ("config/",
     "Các file cấu hình của Laravel (app.php, database.php, mail.php, "
     "filesystems.php, auth.php, shop.php…)."),
    (".env",
     "Biến môi trường nhạy cảm: tên app, APP_KEY, kết nối DB, SMTP, "
     "session driver. Không commit lên Git."),
    ("composer.json / package.json",
     "Khai báo dependencies: composer.json cho gói PHP (Laravel, Breeze, "
     "PhpSpreadsheet…); package.json cho gói NPM (Vite, Bootstrap, Chart.js)."),
    ("vite.config.js",
     "Cấu hình Vite – công cụ bundle CSS/JS frontend."),
    ("artisan",
     "Tiện ích dòng lệnh của Laravel: migrate, db:seed, route:list, "
     "make:controller, serve, storage:link…"),
]

# ----------------------------- helpers -----------------------------
def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def set_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), '000000')
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def format_para(p, *, alignment=None, line_spacing=1.5, size=13,
                bold=False, italic=False, font="Times New Roman",
                first_line_indent=None, color=None):
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.line_spacing = line_spacing
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    for r in p.runs:
        r.font.name = font
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
        if color is not None:
            r.font.color.rgb = color


def insert_paragraph_before(anchor_p, text="", *, style_name=None):
    new_p = OxmlElement('w:p')
    anchor_p._p.addprevious(new_p)
    para = Paragraph(new_p, anchor_p._parent)
    if style_name:
        para.style = anchor_p.part.document.styles[style_name]
    if text:
        para.add_run(text)
    return para


def insert_table_before(anchor_p, rows, cols, doc):
    tbl = doc.add_table(rows=rows, cols=cols)
    tbl.style = "Table Grid"
    anchor_p._p.addprevious(tbl._tbl)
    return tbl


# ----------------------------- main -----------------------------
def main():
    doc = Document(DOC)

    # ---------- locate section 3.1.3 (Heading 3) ----------
    paras = list(doc.paragraphs)
    idx_h313 = None
    for i, p in enumerate(paras):
        if p.style.name == "Heading 3" and p.text.strip().startswith("3.1.3"):
            idx_h313 = i; break
    if idx_h313 is None:
        raise SystemExit("Không tìm thấy heading 3.1.3")
    h313 = paras[idx_h313]

    # ---------- find image paragraph + caption paragraph after heading ----------
    to_delete = []
    for j in range(idx_h313 + 1, len(paras)):
        pj = paras[j]
        if pj.style.name in ("Heading 1", "Heading 2", "Heading 3"):
            break
        to_delete.append(pj)

    next_heading = None
    for j in range(idx_h313 + 1, len(paras)):
        if paras[j].style.name in ("Heading 1", "Heading 2", "Heading 3"):
            next_heading = paras[j]; break
    if next_heading is None:
        raise SystemExit("Không tìm thấy heading kế tiếp sau 3.1.3")

    # ---------- insert intro paragraph ----------
    intro_p = insert_paragraph_before(next_heading, INTRO)
    format_para(intro_p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                first_line_indent=Pt(18))

    # ---------- insert directory-tree code block ----------
    tree_caption_p = insert_paragraph_before(
        next_heading, "Hình 3.1. Cây cấu trúc thư mục dự án Hương Hoa Xinh")
    # Actually no: user wants NO image; the directory tree IS the replacement.
    # We DON'T want to call it 'Hình 3.1' because there's no figure anymore.
    # Remove the caption we just made:
    tree_caption_p._p.getparent().remove(tree_caption_p._p)

    # Create a 1-cell table acting as a code block for the tree
    tree_tbl = insert_table_before(next_heading, rows=1, cols=1, doc=doc)
    cell = tree_tbl.rows[0].cells[0]
    cell.text = ""
    # Set light-grey fill
    set_cell_shading(cell, "F2F2F2")
    set_cell_borders(cell)
    # Each tree line becomes a separate paragraph inside the cell
    lines = TREE.rstrip("\n").split("\n")
    cell.paragraphs[0].text = lines[0]
    for line in lines[1:]:
        cell.add_paragraph(line)
    for cp in cell.paragraphs:
        cp.paragraph_format.line_spacing = 1.15
        cp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        cp.paragraph_format.space_before = Pt(0)
        cp.paragraph_format.space_after = Pt(0)
        cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for r in cp.runs:
            r.font.name = "Consolas"
            r.font.size = Pt(10)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # spacer
    sp = insert_paragraph_before(next_heading, "")
    format_para(sp)

    # ---------- description table heading + table ----------
    desc_intro_p = insert_paragraph_before(
        next_heading,
        "Vai trò chính của các thư mục và file cấu hình quan trọng được tổng "
        "hợp trong bảng dưới đây:")
    format_para(desc_intro_p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                first_line_indent=Pt(18))

    desc_tbl = insert_table_before(next_heading, rows=len(DESC_ROWS) + 1,
                                   cols=2, doc=doc)
    # Header row
    for c, h in enumerate(["Thư mục / File", "Chức năng"]):
        cell = desc_tbl.rows[0].cells[c]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.font.name = "Times New Roman"
        run.font.size = Pt(13)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].paragraph_format.line_spacing = 1.5
        set_cell_shading(cell, "1F4E79")
        set_cell_borders(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # Data rows
    for ri, (path, desc) in enumerate(DESC_ROWS, start=1):
        for ci, val in enumerate([path, desc]):
            cell = desc_tbl.rows[ri].cells[ci]
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            run.font.name = "Consolas" if ci == 0 else "Times New Roman"
            run.font.size = Pt(11 if ci == 0 else 13)
            cell.paragraphs[0].alignment = (
                WD_ALIGN_PARAGRAPH.LEFT if ci == 0
                else WD_ALIGN_PARAGRAPH.JUSTIFY)
            cell.paragraphs[0].paragraph_format.line_spacing = 1.5
            cell.paragraphs[0].paragraph_format.line_spacing_rule = (
                WD_LINE_SPACING.MULTIPLE)
            set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Set first column narrower (about 35%) — done via grid width
    # python-docx makes columns auto by default; this is acceptable.

    # ---------- remove old image + caption paragraphs ----------
    for pdel in to_delete:
        pdel._p.getparent().remove(pdel._p)

    # ============================================================
    # Re-number 'Hình 3.X' in BODY (after section 3.1.3) — X-1
    # ============================================================
    body_pat = re.compile(r'Hình\s*3\.(\d+)')
    # Identify which paragraphs are in danh-mục-hình-ảnh vs body
    # Danh mục is between 'DANH MỤC HÌNH ẢNH' heading and 'DANH MỤC BẢNG BIỂU' heading
    dm_start = dm_end = None
    for i, p in enumerate(doc.paragraphs):
        if p.style.name == "Heading 1" and "HÌNH ẢNH" in p.text.upper():
            dm_start = i
        elif dm_start and p.style.name == "Heading 1":
            dm_end = i; break

    # ---- BODY renumber (paragraphs NOT in [dm_start, dm_end)) ----
    for i, p in enumerate(doc.paragraphs):
        if dm_start is not None and dm_start <= i < dm_end:
            continue
        new_text = body_pat.sub(
            lambda m: f"Hình 3.{int(m.group(1)) - 1}"
            if int(m.group(1)) >= 2 else m.group(0),
            p.text,
        )
        if new_text != p.text:
            # Replace text preserving formatting of first run
            if p.runs:
                # Easiest: dump all runs and put single text on first run
                first = p.runs[0]
                first.text = new_text
                for r in p.runs[1:]:
                    r.text = ""

    # ---- DANH MỤC: remove 'Hình 3.1' entry & renumber rest ----
    if dm_start is not None and dm_end is not None:
        body_paras = doc.paragraphs[dm_start:dm_end]
        to_remove_dm = []
        for p in body_paras:
            m = body_pat.search(p.text)
            if not m:
                continue
            n = int(m.group(1))
            if n == 1:
                to_remove_dm.append(p)
            else:
                new = re.sub(r'Hình\s*3\.\d+',
                             f"Hình 3.{n-1}", p.text, count=1)
                if p.runs:
                    p.runs[0].text = new
                    for r in p.runs[1:]:
                        r.text = ""
        for p in to_remove_dm:
            p._p.getparent().remove(p._p)

        # Add missing 'Hình 3.30. Hồ sơ admin' at the end of section
        # (body has Hình 3.31 → renumbered to Hình 3.30; danh mục doesn't list it).
        # Insert before the next Heading 1 ('DANH MỤC BẢNG BIỂU').
        dm_anchor = doc.paragraphs[dm_end]
        new_entry = insert_paragraph_before(
            dm_anchor, "Hình 3.30. Hồ sơ admin")
        format_para(new_entry, alignment=WD_ALIGN_PARAGRAPH.LEFT)

    # ============================================================
    # Save
    # ============================================================
    doc.save(DOC)
    print("OK")


if __name__ == "__main__":
    main()
