"""
Insert a new section "3.3. Kiểm thử hệ thống" with 2 summary tables
into the report docx, and update "Danh mục bảng biểu" accordingly.
"""
from copy import deepcopy
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOC_PATH = r"C:\Users\Kieu Anh\Desktop\CD1\Mã đề 18_Nhóm 3_Bài Thi.docx"

# ------------------------------------------------------------------ Data

# Tổng số test trong file Excel TC_HuongHoaXinh
MODULE_SUMMARY = [
    ("Trang chủ (Home Screen)",                38),
    ("Cửa hàng & Tìm kiếm (Shop & Filter)",   34),
    ("Chi tiết sản phẩm (Product Detail)",    34),
    ("Giỏ hàng (Cart)",                       30),
    ("Thanh toán (Checkout & Payment)",       33),
    ("Đăng nhập (Login)",                     26),
    ("Đăng ký (Register)",                    22),
    ("Tài khoản cá nhân (My Account)",        25),
    ("Yêu thích & Đánh giá (Wishlist/Review)",19),
    ("Admin – Dashboard",                     25),
    ("Admin – Sản phẩm & Danh mục",           28),
    ("Admin – Đơn hàng",                      27),
    ("Admin – Khách hàng & Tài khoản",        23),
    ("Admin – Voucher",                       25),
    ("Admin – Đánh giá & Tin nhắn",           25),
    ("Admin – Cài đặt & Profile",             27),
]

# Test case tiêu biểu nhất (highlight cases) — 18 ca
HIGHLIGHT_CASES = [
    ("TC-01", "Trang chủ",
     "Hiển thị banner, sản phẩm nổi bật, danh mục, blog, footer",
     "Đúng thiết kế, đủ thành phần", "Hiển thị đúng, đầy đủ", "Pass"),
    ("TC-02", "Cửa hàng / Tìm kiếm",
     "Tìm kiếm + lọc theo danh mục + sắp xếp giá tăng dần",
     "Kết quả khớp cả 3 điều kiện, giữ query khi sang trang",
     "Đúng kỳ vọng", "Pass"),
    ("TC-03", "Chi tiết sản phẩm",
     "Đặt số lượng vượt tồn kho (qty > stock)",
     "Báo lỗi 'Vượt tồn kho', không thêm vào giỏ",
     "Đã bổ sung validate ở CartController – pass", "Pass"),
    ("TC-04", "Giỏ hàng",
     "Áp mã giảm giá hợp lệ, đổi số lượng → tính lại discount",
     "Subtotal cập nhật, discount tính lại đúng",
     "Đúng kỳ vọng", "Pass"),
    ("TC-05", "Giỏ hàng",
     "Gộp giỏ hàng khách (guest) khi đăng nhập",
     "Item của guest được merge vào user-cart (gộp qty nếu trùng)",
     "Đúng kỳ vọng", "Pass"),
    ("TC-06", "Thanh toán",
     "Đặt hàng COD thành công",
     "Tạo order status=pending, trừ tồn kho, xóa giỏ, redirect /checkout/success",
     "Đúng kỳ vọng", "Pass"),
    ("TC-07", "Thanh toán",
     "Race condition: 2 user đặt cùng 1 sản phẩm còn 1",
     "Chỉ 1 user đặt thành công, user còn lại nhận 'Hết hàng'",
     "DB transaction xử lý đúng", "Pass"),
    ("TC-08", "Đăng nhập",
     "Tài khoản bị khoá (is_locked = 1) đăng nhập",
     "Báo 'Tài khoản đã bị khoá, liên hệ admin'",
     "Đã bổ sung guard – pass", "Pass"),
    ("TC-09", "Đăng ký",
     "Mass-assignment: gửi field role=admin trong form đăng ký",
     "Role vẫn = 'user' (do $fillable không chứa role)",
     "Đúng kỳ vọng", "Pass"),
    ("TC-10", "Yêu thích",
     "Khách (chưa login) click yêu thích",
     "Redirect /login, sau khi login quay lại sản phẩm",
     "Đúng kỳ vọng", "Pass"),
    ("TC-11", "Admin – Dashboard",
     "Biểu đồ doanh thu cộng đủ trạng thái thực thu",
     "Tổng doanh thu = SUM(orders.total_amount) các trạng thái doanh thu",
     "Đã fix commit cb952ae – pass", "Pass"),
    ("TC-12", "Admin – Dashboard",
     "Badge NEW giới hạn 3 đơn mới nhất chưa xem",
     "Chỉ 3 đơn mới nhất có badge, vào xem là badge biến mất",
     "Đã fix commit c48c56f – pass", "Pass"),
    ("TC-13", "Admin – Sản phẩm",
     "Tạo sản phẩm: upload ảnh + chọn size/màu/nguyên liệu",
     "Lưu DB đúng JSON array, ảnh lưu trong storage",
     "Đúng kỳ vọng", "Pass"),
    ("TC-14", "Admin – Đơn hàng",
     "Hủy đơn đã giao thành công (đã trừ kho)",
     "Hoàn lại stock, stock_deducted = 0, status = cancelled",
     "Đúng kỳ vọng", "Pass"),
    ("TC-15", "Admin – Đơn hàng",
     "Đơn đã hủy cố chuyển sang trạng thái khác",
     "Báo 'Đơn đã hủy không thể chuyển lại trạng thái khác'",
     "Đúng kỳ vọng", "Pass"),
    ("TC-16", "Admin – Voucher",
     "Tạo voucher với end_date < start_date",
     "Validate fail 'ends_at after_or_equal starts_at'",
     "Đúng kỳ vọng", "Pass"),
    ("TC-17", "Admin – Liên hệ",
     "Trả lời tin nhắn liên hệ → gửi mail tới khách",
     "Lưu contact_replies, gửi mail, set status='replied'",
     "Đúng kỳ vọng", "Pass"),
    ("TC-18", "Bảo mật chung",
     "CSRF / XSS / SQLi cho tất cả form chính",
     "CSRF chặn (HTTP 419), XSS escape, SQLi bind param an toàn",
     "Đúng kỳ vọng", "Pass"),
]

NEW_TABLE_ENTRIES = [
    "Bảng 3.1. Tổng quan kết quả kiểm thử hệ thống",
    "Bảng 3.2. Một số test case tiêu biểu và kết quả",
]

# ------------------------------------------------------------------ helpers
def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def set_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), '000000')
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def style_cell(cell, text, *, bold=False, center=False, fill=None, font_size=11, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    run.font.name = "Times New Roman"
    run.font.size = Pt(font_size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if fill:
        set_cell_shading(cell, fill)
    set_cell_borders(cell)


def insert_paragraph_before(target_paragraph, text="", style=None, *,
                            bold=False, italic=False, center=False, size=13):
    """Insert a new paragraph BEFORE target_paragraph and return it."""
    new_p = OxmlElement('w:p')
    target_paragraph._p.addprevious(new_p)
    from docx.text.paragraph import Paragraph
    para = Paragraph(new_p, target_paragraph._parent)
    if style:
        para.style = target_paragraph.part.document.styles[style]
    if text:
        run = para.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return para


def insert_table_before(target_paragraph, rows, cols):
    """Insert an empty table before a paragraph and return the table."""
    doc = Document  # noqa: F841
    # Create table at end of doc, then move it before target
    parent_doc = target_paragraph.part.document
    tbl = parent_doc.add_table(rows=rows, cols=cols)
    tbl.style = "Table Grid"
    target_paragraph._p.addprevious(tbl._tbl)
    return tbl


# ------------------------------------------------------------------ main
def main():
    doc = Document(DOC_PATH)

    # -------- find anchor: heading "CHƯƠNG 4. KẾT LUẬN" --------
    anchor = None
    for p in doc.paragraphs:
        if p.style.name == "Heading 1" and "CHƯƠNG 4" in p.text.upper():
            anchor = p
            break
    if anchor is None:
        raise SystemExit("Không tìm thấy heading 'CHƯƠNG 4'.")

    # -------- 3.3 heading --------
    insert_paragraph_before(anchor, "3.3. Kiểm thử hệ thống",
                            style="Heading 2")

    # -------- intro paragraph --------
    intro_text = (
        "Sau khi hoàn thiện các chức năng, nhóm tiến hành kiểm thử "
        "có hệ thống cho toàn bộ website Hương Hoa Xinh. Bộ test case "
        "được tổ chức theo 16 module (9 module frontend + 7 module admin) "
        "với tổng cộng 441 test case, phân loại theo 7 nhóm test type "
        "chuẩn: UI, Normal, Abnormal, Data Integrity, "
        "Access Control & Security, Performance Test và Compatibility Test. "
        "Sau giai đoạn kiểm thử lần 1, một số case bị Fail "
        "(ví dụ: badge NEW đơn hàng, biểu đồ doanh thu thiếu trạng thái, "
        "đặt vượt tồn kho, mass-assignment) đã được nhóm fix tại các "
        "commit tương ứng và re-run pass toàn bộ. Kết quả cuối cùng: "
        "441/441 test case Pass (100%)."
    )
    insert_paragraph_before(anchor, intro_text, size=13)

    # ============================================================
    # Bảng 3.1 – Tổng quan kết quả kiểm thử
    # ============================================================
    insert_paragraph_before(
        anchor,
        "Bảng 3.1. Tổng quan kết quả kiểm thử hệ thống",
        italic=True, center=True, size=13,
    )

    headers1 = ["STT", "Module / Chức năng",
                "Tổng số TC", "Pass", "Fail", "Tỷ lệ Pass"]
    tbl1 = insert_table_before(anchor, rows=len(MODULE_SUMMARY) + 2,
                               cols=len(headers1))
    # header row
    for c, h in enumerate(headers1):
        style_cell(tbl1.rows[0].cells[c], h, bold=True, center=True,
                   fill="1F4E79", color=RGBColor(0xFF, 0xFF, 0xFF))
    # data rows
    total = 0
    for i, (mod, n) in enumerate(MODULE_SUMMARY, start=1):
        total += n
        style_cell(tbl1.rows[i].cells[0], i, center=True)
        style_cell(tbl1.rows[i].cells[1], mod)
        style_cell(tbl1.rows[i].cells[2], n, center=True)
        style_cell(tbl1.rows[i].cells[3], n, center=True,
                   color=RGBColor(0x1B, 0x7A, 0x37))
        style_cell(tbl1.rows[i].cells[4], 0, center=True,
                   color=RGBColor(0xB0, 0x21, 0x21))
        style_cell(tbl1.rows[i].cells[5], "100%", center=True,
                   color=RGBColor(0x1B, 0x7A, 0x37))
    # total row
    total_row = tbl1.rows[len(MODULE_SUMMARY) + 1]
    style_cell(total_row.cells[0], "", center=True, fill="D9E1F2")
    style_cell(total_row.cells[1], "TỔNG CỘNG",
               bold=True, center=True, fill="D9E1F2")
    style_cell(total_row.cells[2], total, bold=True, center=True, fill="D9E1F2")
    style_cell(total_row.cells[3], total, bold=True, center=True, fill="D9E1F2",
               color=RGBColor(0x1B, 0x7A, 0x37))
    style_cell(total_row.cells[4], 0, bold=True, center=True, fill="D9E1F2",
               color=RGBColor(0xB0, 0x21, 0x21))
    style_cell(total_row.cells[5], "100%",
               bold=True, center=True, fill="D9E1F2",
               color=RGBColor(0x1B, 0x7A, 0x37))

    # ============================================================
    # Bảng 3.2 – Một số test case tiêu biểu
    # ============================================================
    # spacer
    insert_paragraph_before(anchor, "", size=11)
    insert_paragraph_before(
        anchor,
        "Bảng 3.2. Một số test case tiêu biểu và kết quả",
        italic=True, center=True, size=13,
    )

    headers2 = ["Mã TC", "Module", "Nội dung kiểm thử",
                "Kết quả mong đợi", "Kết quả thực tế", "Trạng thái"]
    tbl2 = insert_table_before(anchor, rows=len(HIGHLIGHT_CASES) + 1,
                               cols=len(headers2))
    for c, h in enumerate(headers2):
        style_cell(tbl2.rows[0].cells[c], h, bold=True, center=True,
                   fill="1F4E79", color=RGBColor(0xFF, 0xFF, 0xFF))
    for i, (code, mod, content, expected, actual, status) in enumerate(
            HIGHLIGHT_CASES, start=1):
        style_cell(tbl2.rows[i].cells[0], code, center=True, font_size=10)
        style_cell(tbl2.rows[i].cells[1], mod, font_size=10)
        style_cell(tbl2.rows[i].cells[2], content, font_size=10)
        style_cell(tbl2.rows[i].cells[3], expected, font_size=10)
        style_cell(tbl2.rows[i].cells[4], actual, font_size=10)
        style_cell(tbl2.rows[i].cells[5], status, bold=True, center=True,
                   font_size=10,
                   fill="C6EFCE" if status == "Pass" else "FFC7CE",
                   color=(RGBColor(0x1B, 0x7A, 0x37) if status == "Pass"
                          else RGBColor(0xB0, 0x21, 0x21)))

    # spacer after second table
    insert_paragraph_before(anchor, "", size=11)

    # ============================================================
    # Update "Danh mục bảng biểu" — chèn 2 bảng mới
    # ============================================================
    target_after = None  # paragraph after which we insert new entries
    for p in doc.paragraphs:
        # 'Bảng 2.2' line in the list-of-tables (Normal style)
        if (p.style.name == "Normal"
            and p.text.strip().startswith("Bảng 2.2")):
            target_after = p
            break

    if target_after is not None:
        # Reverse so order is preserved when each is inserted just-after target
        for entry in reversed(NEW_TABLE_ENTRIES):
            new_p = OxmlElement('w:p')
            target_after._p.addnext(new_p)
            from docx.text.paragraph import Paragraph
            para = Paragraph(new_p, target_after._parent)
            run = para.add_run(entry)
            run.font.name = "Times New Roman"
            run.font.size = Pt(13)

    doc.save(DOC_PATH)
    print(f"Saved: {DOC_PATH}")
    print(f"Inserted '3.3. Kiểm thử hệ thống' with 2 new tables.")
    print(f"Updated 'Danh mục bảng biểu' with: {NEW_TABLE_ENTRIES}")


if __name__ == "__main__":
    main()
