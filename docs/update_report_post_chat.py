"""
Update analysis sections of the report after migrating contact-form
to direct chat + checkout sender/recipient flow.
"""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING

DOC = r"C:\Users\Kieu Anh\Desktop\CD1\Mã đề 18_Nhóm 3_Bài Thi.docx"


def replace_paragraph_text(p, new_text, *, bold_prefix=None):
    """Replace whole paragraph text, preserving font/size of first run."""
    if not p.runs:
        run = p.add_run(new_text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(13)
        return
    first = p.runs[0]
    # Wipe extra runs
    for r in p.runs[1:]:
        r.text = ""
    first.text = ""
    if bold_prefix:
        rb = p.add_run(bold_prefix)
        rb.font.name = "Times New Roman"
        rb.font.size = Pt(13)
        rb.bold = True
        rn = p.add_run(new_text)
        rn.font.name = "Times New Roman"
        rn.font.size = Pt(13)
    else:
        first.text = new_text
        first.font.name = "Times New Roman"
        first.font.size = Pt(13)


def replace_cell_text(cell, new_text):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(new_text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(13)
    p.paragraph_format.line_spacing = 1.5


def find_paragraph_starting_with(doc, prefix, start_index=0):
    for i, p in enumerate(doc.paragraphs):
        if i < start_index: continue
        if p.text.strip().startswith(prefix):
            return i, p
    return None, None


def main():
    doc = Document(DOC)

    # ============================================================
    # 1. Section 1.1 — Liên hệ bullet → Chat
    # ============================================================
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith("Liên hệ:") and "tin nhắn liên hệ" in t:
            replace_paragraph_text(p,
                "Chat trực tiếp: Khách hàng đã đăng nhập trao đổi 2 chiều "
                "với cửa hàng qua hộp thoại chat (lưu DB, không qua email); "
                "có nút chat nổi ở mọi trang kèm badge thông báo tin mới; "
                "admin phản hồi ngay trong giao diện quản trị.",
                bold_prefix="Chat trực tiếp: ".replace("Chat trực tiếp: ", "")  # we'll prepend below
            )
            # The above call kept it simple; re-bold prefix manually
            # Better just keep as plain
            break

    # Re-do this paragraph cleanly without bold prefix (Word style usually
    # uses plain text in list paragraphs)
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith("Chat trực tiếp:"):
            # Already updated above
            break

    # ============================================================
    # 2. Section 1.1 — Giỏ hàng & Thanh toán bullet (mention sender/recipient)
    # ============================================================
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith("Giỏ hàng & Thanh toán:"):
            replace_paragraph_text(p,
                "Giỏ hàng & Thanh toán: Thêm/sửa/xoá sản phẩm trong giỏ; "
                "Guest Cart cho khách vãng lai; áp voucher; tính phí ship "
                "theo khoảng cách; tách thông tin người gửi và người nhận, "
                "chọn ngày + khung giờ giao, viết lời nhắn in trên thiệp "
                "kèm hoa; thanh toán COD hoặc chuyển khoản QR.")
            break

    # ============================================================
    # 3. Section 1.2 — US-14 in Table 1
    # ============================================================
    user_story_table = doc.tables[1]
    for row in user_story_table.rows:
        cells = row.cells
        if cells[0].text.strip() == "US-14":
            replace_cell_text(cells[2],
                "Là admin, tôi muốn chat 2 chiều trực tiếp với khách hàng "
                "(lưu DB, hỗ trợ realtime) để xem lịch sử trò chuyện và phản "
                "hồi nhanh trong giao diện quản trị.")
            replace_cell_text(cells[3], "Cao")
            break

    # Add new US-17 row for "Gửi hoa cho người khác"
    has_us17 = any(row.cells[0].text.strip() == "US-17" for row in user_story_table.rows)
    if not has_us17:
        new_row = user_story_table.add_row()
        replace_cell_text(new_row.cells[0], "US-17")
        replace_cell_text(new_row.cells[1], "Customer")
        replace_cell_text(new_row.cells[2],
            "Là khách hàng, tôi muốn gửi hoa cho người khác — nhập tên / "
            "SĐT / địa chỉ người nhận, chọn ngày & khung giờ giao, kèm lời "
            "nhắn in trên thiệp.")
        replace_cell_text(new_row.cells[3], "Cao")

    # ============================================================
    # 4. Section 1.3 — Table 2 (tác nhân)
    # ============================================================
    actors_table = doc.tables[2]
    for row in actors_table.rows:
        c0 = row.cells[0].text.strip()
        # Guest row
        if "Khách vãng lai" in row.cells[1].text:
            replace_cell_text(row.cells[4],
                "Xem sản phẩm, tìm kiếm, đọc blog, đăng ký / đăng nhập.")
        elif "Khách hàng" in row.cells[1].text:
            replace_cell_text(row.cells[4],
                "Mua hàng, quản lý giỏ, thanh toán (tách người gửi & người "
                "nhận), wishlist, viết đánh giá, chat trực tiếp với cửa hàng.")
        elif "Quản trị" in row.cells[1].text:
            replace_cell_text(row.cells[4],
                "CRUD sản phẩm, danh mục, voucher; xử lý đơn; quản lý khách "
                "hàng; chat trả lời khách; xem báo cáo doanh thu.")
        elif "Hệ thống" in row.cells[1].text:
            replace_cell_text(row.cells[4],
                "Tự huỷ đơn pending quá 10 phút và hoàn tồn kho; polling "
                "tin nhắn chat 4 giây/lần cho 2 phía.")
        elif "Email" in row.cells[1].text:
            # Email service still used for password reset only
            replace_cell_text(row.cells[3],
                "Dịch vụ SMTP bên ngoài (chỉ dùng cho xác thực tài khoản).")
            replace_cell_text(row.cells[4],
                "Gửi email reset mật khẩu, xác thực email khi đăng ký. "
                "Không còn dùng cho luồng liên hệ — đã thay bằng chat DB.")

    # ============================================================
    # 5. Section 2.3 ERD — Table 3 (add new relationships)
    # ============================================================
    rel_table = doc.tables[3]
    has_user_contact = any(
        ("contact_messages" in row.cells[2].text and "users" in row.cells[1].text)
        or ("contact_messages" in row.cells[1].text and "users" in row.cells[2].text)
        for row in rel_table.rows
    )
    if not has_user_contact:
        # Insert 2 new rows
        new_idx = len(rel_table.rows)
        row1 = rel_table.add_row()
        replace_cell_text(row1.cells[0], str(new_idx))
        replace_cell_text(row1.cells[1], "users")
        replace_cell_text(row1.cells[2], "contact_messages")
        replace_cell_text(row1.cells[3], "1 – 1 (mềm)")
        replace_cell_text(row1.cells[4],
            "Mỗi khách hàng có một hộp thoại chat chính với shop (qua user_id).")

        new_idx += 1
        row2 = rel_table.add_row()
        replace_cell_text(row2.cells[0], str(new_idx))
        replace_cell_text(row2.cells[1], "users")
        replace_cell_text(row2.cells[2], "contact_replies")
        replace_cell_text(row2.cells[3], "1 – N")
        replace_cell_text(row2.cells[4],
            "Khách hàng cũng có thể là người gửi reply trong chat "
            "(qua user_id, song song với admin_id của bên shop).")

    # ============================================================
    # 6. Section 2.3 — note paragraph after Bảng 2.2 (idx ~189)
    # ============================================================
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t.startswith("Ngoài ra, bảng categories có quan hệ tự tham chiếu"):
            # Append updated info to this paragraph
            new_text = (
                "Ngoài ra, bảng categories có quan hệ tự tham chiếu "
                "(parent_id → categories.id) cho phép biểu diễn danh mục "
                "cha – danh mục con. Mối quan hệ giữa vouchers và orders "
                "là quan hệ \"lỏng\": không có cột voucher_id trong "
                "orders, chỉ lưu mã voucher trong note để truy vết. "
                "Phiên bản hiện tại bổ sung thêm các cột: contact_messages."
                "user_id (1 user ↔ 1 hộp thoại chat), contact_replies."
                "user_id và customer_read_at (khách phản hồi & badge đã đọc), "
                "và 6 cột vào orders cho luồng tách người gửi / người nhận "
                "(sender_name, sender_phone, recipient_name, delivery_date, "
                "delivery_time_slot, recipient_message)."
            )
            replace_paragraph_text(p, new_text)
            break

    # ============================================================
    # 7. Wireframe captions (Hình 2.21)
    # ============================================================
    for p in doc.paragraphs:
        if p.text.strip() == "Hình 2.21. Wireframe trang Liên hệ":
            replace_paragraph_text(p, "Hình 2.21. Wireframe trang Chat – Hộp thoại trực tiếp với cửa hàng")
            break

    # ============================================================
    # 8. Wireframe captions (Hình 3.7)
    # ============================================================
    for p in doc.paragraphs:
        if "Hình 3.7." in p.text and "Liên hệ" in p.text:
            replace_paragraph_text(p,
                "Hình 3.7. Trang Chat – Hộp thoại trực tiếp với cửa hàng (frontend)")
            break

    # ============================================================
    # 9. Wireframe captions (Hình 3.27 = Admin tin nhắn)
    # ============================================================
    for p in doc.paragraphs:
        if "Hình 3.27." in p.text and "in nhắn" in p.text:
            replace_paragraph_text(p,
                "Hình 3.27. Admin – Chat với khách hàng (hộp thoại realtime)")
            break

    # ============================================================
    # 10. Section 4.1 — kết quả
    # ============================================================
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith("•  Xây dựng đầy đủ khu vực admin"):
            replace_paragraph_text(p,
                "•  Xây dựng đầy đủ khu vực admin với guard riêng: "
                "dashboard, quản lý sản phẩm / danh mục / đơn hàng / "
                "khách hàng / voucher / đánh giá / chat với khách hàng / "
                "cài đặt website.")
            break

    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith("•  Tích hợp gửi email tự động"):
            replace_paragraph_text(p,
                "•  Triển khai hệ thống chat 2 chiều khách – admin lưu "
                "trực tiếp trong DB (không qua email), polling realtime mỗi "
                "4 giây, có nút chat nổi + badge thông báo tin mới ở mọi "
                "trang khách hàng; tự huỷ đơn pending quá 10 phút và hoàn "
                "tồn kho; email chỉ dùng cho xác thực và reset mật khẩu.")
            break

    # Thêm bullet về tính năng người gửi/người nhận
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t.startswith("•  Áp dụng đúng mô hình MVC"):
            # Insert a new bullet just before this paragraph
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            from docx.text.paragraph import Paragraph
            new_p = OxmlElement('w:p')
            p._p.addprevious(new_p)
            np = Paragraph(new_p, p._parent)
            run = np.add_run(
                "•  Tách rõ thông tin người gửi và người nhận trong luồng "
                "thanh toán: cho phép đặt hoa tặng người khác với ngày + "
                "khung giờ giao và lời nhắn in trên thiệp; hiển thị đầy đủ "
                "trong trang xác nhận đơn và trang chi tiết đơn của admin."
            )
            run.font.name = "Times New Roman"
            run.font.size = Pt(13)
            np.paragraph_format.line_spacing = 1.5
            np.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            break

    # ============================================================
    # 11. Test case Bảng 3.2 — TC-17 update
    # ============================================================
    for t in doc.tables:
        if not t.rows: continue
        head = [c.text.strip() for c in t.rows[0].cells]
        if head[:2] == ['Mã TC', 'Module']:
            for row in t.rows[1:]:
                if row.cells[0].text.strip() == "TC-17":
                    replace_cell_text(row.cells[1], "Admin – Chat")
                    replace_cell_text(row.cells[2],
                        "Admin trả lời tin nhắn của khách trong hộp thoại chat")
                    replace_cell_text(row.cells[3],
                        "Lưu contact_replies với admin_id, status='replied'; "
                        "khách thấy tin mới qua polling 4s + badge bật")
                    replace_cell_text(row.cells[4], "Đúng kỳ vọng")
                    replace_cell_text(row.cells[5], "Pass")
                    break
            break

    doc.save(DOC)
    print("Updated docx successfully")


if __name__ == "__main__":
    main()
