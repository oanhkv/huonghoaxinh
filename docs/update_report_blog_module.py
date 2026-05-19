"""
Update the report .docx after the Blog admin module was added:
  - 3.1.2: bổ sung BlogPostController/BlogCategoryController vào nhánh Admin
  - 3.1.4: thêm subsection (d) — code sample BlogPostController + Quill editor
  - 3.2.2: thêm 4 figure captions Hình 3.32–3.35
  - Danh mục hình ảnh: thêm 4 entries tương ứng
  - 4.1: thêm bullet về Blog CMS với rich-text editor
"""
import re, sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from copy import deepcopy
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

DOC = r"C:\Users\Kieu Anh\Desktop\CD1\Mã đề 18_Nhóm 3_Bài Thi.docx"


# ----- helpers -----
def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def set_cell_borders(cell, color='808080'):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_b = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0');    b.set(qn('w:color'), color)
        tc_b.append(b)
    tc_pr.append(tc_b)


def new_para_before(anchor, segments=None, *, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                    line_spacing=1.5, first_line_indent=None, style_name=None,
                    font="Times New Roman", size_pt=13):
    new_p = OxmlElement('w:p')
    anchor._p.addprevious(new_p)
    p = Paragraph(new_p, anchor._parent)
    if style_name:
        p.style = anchor.part.document.styles[style_name]
    if segments is None:
        segments = [("", False, False)]
    if isinstance(segments, str):
        segments = [(segments, False, False)]
    for seg in segments:
        if len(seg) == 2: txt, bold = seg; italic = False
        else: txt, bold, italic = seg
        run = p.add_run(txt)
        run.font.name = font
        run.font.size = Pt(size_pt)
        run.bold = bold
        run.italic = italic
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.line_spacing = line_spacing
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    return p


def new_code_block_before(anchor, code_text, doc):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    anchor._p.addprevious(tbl._tbl)
    cell = tbl.rows[0].cells[0]
    cell.text = ""
    set_cell_shading(cell, "F5F5F5")
    set_cell_borders(cell)
    lines = code_text.rstrip("\n").split("\n")
    cell.paragraphs[0].text = lines[0]
    for line in lines[1:]:
        cell.add_paragraph(line)
    for cp in cell.paragraphs:
        cp.paragraph_format.line_spacing = 1.15
        cp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        cp.paragraph_format.space_before = Pt(0)
        cp.paragraph_format.space_after = Pt(0)
        cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for r in cp.runs:
            r.font.name = "Consolas"
            r.font.size = Pt(10)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    return tbl


def patch_paragraph_text(p, new_text):
    if not p.runs:
        run = p.add_run(new_text)
        run.font.name = "Times New Roman"; run.font.size = Pt(13)
        return
    first = p.runs[0]
    first.text = new_text
    first.font.name = "Times New Roman"; first.font.size = Pt(13)
    for r in p.runs[1:]:
        r.text = ""


# ----------------------------- 1. Section 3.1.2 — update Controllers
def update_312(doc):
    """Add BlogPostController/BlogCategoryController mention to Controllers bullet."""
    for p in doc.paragraphs:
        t = p.text.strip()
        if "Admin/" in t and "DashboardController" in t and "Controller cho khu vực quản trị" in t:
            # Found the Admin controllers bullet — patch it to include Blog controllers
            # Preserve italic runs; simplest: rewrite plain
            new_text = (
                "•  Admin/ – chứa Controller cho khu vực quản trị: DashboardController, ProductController, "
                "CategoryController, OrderController, UserController, VoucherController, ReviewController, "
                "RevenueController, ContactMessageController, WebsiteSettingController, AdminProfileController, "
                "BlogPostController, BlogCategoryController. Mỗi Controller phụ trách một module CRUD riêng."
            )
            patch_paragraph_text(p, new_text)
            print("  3.1.2 Admin Controllers bullet — updated")
            return


# ----------------------------- 2. Section 3.1.4 — add (d) Blog module code
def update_314(doc):
    """Insert subsection (d) Blog admin with Quill code before heading 3.1.5."""
    h315 = None
    for p in doc.paragraphs:
        if p.style.name == "Heading 3" and p.text.strip().startswith("3.1.5"):
            h315 = p; break
    if h315 is None:
        print("  Không tìm thấy heading 3.1.5"); return

    # Insert: heading 4 + paragraph + code block (controller) + paragraph + code block (Quill)
    h4 = new_para_before(h315, "d) Blog Admin – BlogPostController + Quill Editor")
    h4.style = doc.styles["Heading 4"]

    new_para_before(h315,
        "Module Blog quản trị cho phép admin đăng / sửa / xoá bài viết, dùng "
        "Quill 2.0 làm trình soạn thảo rich text. Controller bên dưới xử lý "
        "luồng `index` với tìm kiếm + lọc theo danh mục + lọc trạng thái:",
        first_line_indent=Pt(18))

    new_code_block_before(h315,
        "// app/Http/Controllers/Admin/BlogPostController.php\n"
        "public function index(Request $request)\n"
        "{\n"
        "    $query = BlogPost::query()->with('category')->latest('id');\n"
        "\n"
        "    if ($request->filled('q')) {\n"
        "        $q = $request->q;\n"
        "        $query->where(function ($sub) use ($q) {\n"
        "            $sub->where('title', 'like', \"%{$q}%\")\n"
        "                ->orWhere('slug', 'like', \"%{$q}%\")\n"
        "                ->orWhere('excerpt', 'like', \"%{$q}%\");\n"
        "        });\n"
        "    }\n"
        "    if ($request->filled('category')) {\n"
        "        $query->where('blog_category_id', $request->category);\n"
        "    }\n"
        "    if ($request->filled('status')) {\n"
        "        $query->where('is_active', $request->status === 'active');\n"
        "    }\n"
        "\n"
        "    $posts      = $query->paginate(10)->withQueryString();\n"
        "    $categories = BlogCategory::withCount('posts')\n"
        "                    ->orderByDesc('posts_count')->get();\n"
        "    $stats = [\n"
        "        'total'      => BlogPost::count(),\n"
        "        'active'     => BlogPost::where('is_active', true)->count(),\n"
        "        'inactive'   => BlogPost::where('is_active', false)->count(),\n"
        "        'categories' => BlogCategory::count(),\n"
        "    ];\n"
        "    return view('admin.blog_posts.index',\n"
        "        compact('posts', 'categories', 'stats'));\n"
        "}",
        doc=doc)

    new_para_before(h315,
        "Trình soạn thảo Quill được khởi tạo bằng JavaScript với 2 toolbar — "
        "toolbar nhỏ cho 'Tóm tắt' và toolbar đầy đủ cho 'Nội dung chính' "
        "(heading, font, cỡ chữ, B/I/U/S, màu chữ, căn lề, danh sách, indent, "
        "blockquote, code, link, ảnh, clear):",
        first_line_indent=Pt(18))

    new_code_block_before(h315,
        "// resources/views/admin/blog_posts/form.blade.php (đoạn JS)\n"
        "const mainToolbar = [\n"
        "    [{ header: [1,2,3,4,false] }, { font: [] }],\n"
        "    [{ size: ['small', false, 'large', 'huge'] }],\n"
        "    ['bold', 'italic', 'underline', 'strike'],\n"
        "    [{ color: [] }, { background: [] }],\n"
        "    [{ align: [] }, { list: 'ordered' }, { list: 'bullet' }],\n"
        "    [{ indent: '-1' }, { indent: '+1' }],\n"
        "    ['blockquote', 'code-block', 'link', 'image'],\n"
        "    ['clean'],\n"
        "];\n"
        "\n"
        "const contentQuill = new Quill('#contentEditor', {\n"
        "    theme: 'snow',\n"
        "    placeholder: 'Bắt đầu viết bài tại đây…',\n"
        "    modules: { toolbar: mainToolbar }\n"
        "});\n"
        "\n"
        "// Trước khi submit: kéo HTML từ editor về hidden input\n"
        "form.addEventListener('submit', function () {\n"
        "    contentHidden.value = contentQuill.root.innerHTML.trim();\n"
        "    excerptHidden.value = excerptQuill.root.innerHTML.trim();\n"
        "});",
        doc=doc)
    print("  3.1.4 (d) Blog module — đã thêm 2 code block + heading 4")


# ----------------------------- 3. Section 3.2.2 — add 4 new captions
def update_322(doc):
    """Insert 4 figure captions Hình 3.32–3.35 before heading Chương 4."""
    h_ch4 = None
    for p in doc.paragraphs:
        if p.style.name == "Heading 1" and "CHƯƠNG 4" in p.text.upper():
            h_ch4 = p; break
    if h_ch4 is None:
        print("  Không tìm thấy Chương 4"); return

    captions = [
        "Hình 3.32. Admin – Quản lý bài viết Blog (stats cards + filter pills + grid 2 cột)",
        "Hình 3.33. Admin – Form tạo / sửa bài viết với Quill rich text editor (2 toolbar đầy đủ)",
        "Hình 3.34. Admin – Danh mục blog (bảng + nút tạo mới + search)",
        "Hình 3.35. Admin – Form tạo / sửa danh mục blog",
    ]
    # Insert intro paragraph + captions before Chương 4
    intro = new_para_before(h_ch4,
        "Module Blog ở khu vực quản trị bổ sung 4 màn hình mới cho phép admin "
        "đăng / sửa / xoá bài viết và phân loại theo danh mục hoa — sử dụng "
        "Quill Editor cho phần soạn thảo nội dung:",
        first_line_indent=Pt(18))

    for cap in captions:
        p = new_para_before(h_ch4, cap)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.italic = True
        # spacer line after each
        new_para_before(h_ch4, "")
    print(f"  3.2.2 — đã thêm {len(captions)} caption Hình 3.32–3.35")


# ----------------------------- 4. Danh mục hình ảnh — add 4 new entries
def update_danh_muc(doc):
    """Insert 4 new figure entries after Hình 3.31 in the figure list."""
    # Find the Hinh 3.31 entry in the figure-list area (within first 130 paragraphs)
    target = None
    for i, p in enumerate(doc.paragraphs):
        if i > 130: break
        t = p.text.strip()
        if t.startswith("Hình 3.31"):
            target = p
    if target is None:
        print("  Không tìm thấy Hình 3.31 trong danh mục hình ảnh"); return

    captions = [
        "Hình 3.32. Admin – Quản lý bài viết Blog",
        "Hình 3.33. Admin – Form tạo / sửa bài viết (Quill editor)",
        "Hình 3.34. Admin – Danh mục blog",
        "Hình 3.35. Admin – Form tạo / sửa danh mục blog",
    ]
    # Insert after target — use addnext, but in reverse so they appear in order
    parent = target._p.getparent()
    for cap in reversed(captions):
        new_p = OxmlElement('w:p')
        target._p.addnext(new_p)
        para = Paragraph(new_p, target._parent)
        run = para.add_run(cap)
        run.font.name = "Times New Roman"; run.font.size = Pt(13)
        para.paragraph_format.line_spacing = 1.5
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    print(f"  Danh mục hình ảnh — đã thêm {len(captions)} entry Hình 3.32–3.35")


# ----------------------------- 5. Section 4.1 — add Blog CMS bullet
def update_41(doc):
    """Add a new bullet about Blog CMS after the existing chat bullet."""
    # Find the chat bullet to insert after
    target = None
    for p in doc.paragraphs:
        if p.text.strip().startswith("•  Triển khai hệ thống chat 2 chiều"):
            target = p; break
    if target is None:
        print("  Không tìm thấy chat bullet ở 4.1"); return

    # Insert a new bullet AFTER target
    new_p = OxmlElement('w:p')
    target._p.addnext(new_p)
    para = Paragraph(new_p, target._parent)
    run = para.add_run(
        "•  Bổ sung Blog CMS đầy đủ ở khu vực admin: CRUD bài viết, CRUD danh "
        "mục, lọc đa tiêu chí (search + danh mục + trạng thái), thống kê nhanh "
        "(4 stat cards), pills lọc danh mục. Trình soạn thảo Quill 2.0 cho "
        "phép thay đổi cỡ chữ – font – căn lề – màu chữ – danh sách – ảnh – "
        "blockquote – code, và lưu đúng format khi render ở trang Blog."
    )
    run.font.name = "Times New Roman"; run.font.size = Pt(13)
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    print("  4.1 — đã thêm bullet Blog CMS")


def main():
    doc = Document(DOC)
    update_312(doc)
    update_314(doc)
    update_322(doc)
    update_danh_muc(doc)
    update_41(doc)
    doc.save(DOC)
    print("OK")


if __name__ == "__main__":
    main()
