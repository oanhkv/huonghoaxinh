"""
1. Xóa heading "2.3.1" + các đoạn giải thích cũ đã thêm
2. Thêm 3 caption ảnh ERD (Hình 2.5/2.6/2.7) + giải thích ngay sau từng ảnh trong section 2.3
3. Renumber: tất cả "Hình 2.X" trong body và Danh mục hình ảnh với X ≥ 5 → X+3
"""
import sys, io, re
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

DOC = r"C:\Users\Kieu Anh\Desktop\CD1\Mã đề 18_Nhóm 3_Bài Thi.docx"


def new_para_before(anchor, segments, *,
                    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                    line_spacing=1.5,
                    first_line_indent=None,
                    italic_default=False,
                    bold_default=False,
                    font="Times New Roman", size_pt=13):
    new_p = OxmlElement('w:p')
    anchor._p.addprevious(new_p)
    p = Paragraph(new_p, anchor._parent)
    if isinstance(segments, str):
        segments = [(segments, bold_default, italic_default)]
    for seg in segments:
        if len(seg) == 2: txt, bold = seg; italic = italic_default
        else: txt, bold, italic = seg
        run = p.add_run(txt)
        run.font.name = font
        run.font.size = Pt(size_pt)
        run.bold = bold
        run.italic = italic
    p.alignment = alignment
    pf = p.paragraph_format
    pf.line_spacing = line_spacing
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    return p


# ============================================================
# 1. Xoá heading 2.3.1 và các đoạn giải thích cũ
# ============================================================
def remove_old_subsection(doc):
    to_remove = []
    paras = list(doc.paragraphs)
    for i, p in enumerate(paras):
        t = p.text.strip()
        if (t == "2.3.1. Lý do chia ERD thành 3 sơ đồ và mối liên hệ"
            or t.startswith("Do hệ thống có tới 16 bảng")
            or t == "Cụ thể, ba sơ đồ ERD được phân chia như sau:"
            or t.startswith("•  Sơ đồ ERD 1 — Khách hàng")
            or t.startswith("•  Sơ đồ ERD 2 — Đơn hàng")
            or t.startswith("•  Sơ đồ ERD 3 — Nội dung")
            or t.startswith("Mối liên kết giữa ba sơ đồ")
            or t.startswith("Cách chia này tuân theo nguyên tắc")):
            to_remove.append(p)
    for p in to_remove:
        p._p.getparent().remove(p._p)
    print(f"  Đã xoá {len(to_remove)} đoạn cũ của mục 2.3.1")


# ============================================================
# 2. Renumber Hình 2.X (X ≥ 5) → +3 ở body + danh mục
# ============================================================
HINH_PAT = re.compile(r'^Hình\s+2\.(\d+)(.*)$')


def renumber_caption(text, shift=3):
    m = HINH_PAT.match(text.strip())
    if not m:
        return None
    num = int(m.group(1))
    rest = m.group(2)
    if num < 5:
        return None
    return f"Hình 2.{num + shift}{rest}"


def renumber_all(doc):
    count = 0
    for p in doc.paragraphs:
        t = p.text.strip()
        new = renumber_caption(t, shift=3)
        if new is not None:
            # Replace text preserving first run style
            full = p.text
            new_full = full.replace(t, new) if t in full else new
            # Replace only the FIRST run content with new text
            if p.runs:
                p.runs[0].text = new_full
                for r in p.runs[1:]:
                    r.text = ""
            count += 1
    print(f"  Đã renumber {count} caption Hình 2.X → +3")


# ============================================================
# 3. Thêm 3 caption + giải thích vào section 2.3
# ============================================================
def add_erd_captions(doc):
    # Find heading 2.4 to insert before
    anchor = None
    for p in doc.paragraphs:
        if p.style.name == "Heading 2" and p.text.strip().startswith("2.4"):
            anchor = p; break
    if anchor is None:
        print("  Không tìm thấy heading 2.4"); return

    # Idempotency: nếu đã có Hình 2.5 = "Sơ đồ ERD 1" rồi thì bỏ qua
    for p in doc.paragraphs:
        if "Hình 2.5. Sơ đồ ERD 1" in p.text:
            print("  Đã có Hình 2.5 ERD — bỏ qua"); return

    # ===== Intro paragraph =====
    new_para_before(anchor,
        "Do hệ thống có 16 bảng với mật độ quan hệ dày đặc (riêng bảng users "
        "có 7 mối quan hệ trực tiếp), nhóm em chia sơ đồ ERD thành 3 sơ đồ "
        "con theo từng nhóm chức năng để dễ trình bày và đọc hiểu — vẫn đảm "
        "bảo đầy đủ thông tin của cả 16 bảng:",
        first_line_indent=Pt(18))

    # ===== Hình 2.5 — ERD 1 =====
    new_para_before(anchor, "[ Chèn ảnh sơ đồ ERD 1 ở đây ]",
                    alignment=WD_ALIGN_PARAGRAPH.CENTER, italic_default=True)
    p = new_para_before(anchor,
        "Hình 2.5. Sơ đồ ERD 1 — Khách hàng & Sản phẩm",
        alignment=WD_ALIGN_PARAGRAPH.CENTER, italic_default=True)
    new_para_before(anchor,
        "Sơ đồ ERD 1 gồm 6 thực thể chính: users, categories, products, "
        "carts, wishlists, reviews — đây là phần lõi của module bán hàng. "
        "Mọi luồng nghiệp vụ khác đều bắt đầu từ hai bảng users hoặc "
        "products được vẽ chi tiết tại đây.",
        first_line_indent=Pt(18))

    # ===== Hình 2.6 — ERD 2 =====
    new_para_before(anchor, "[ Chèn ảnh sơ đồ ERD 2 ở đây ]",
                    alignment=WD_ALIGN_PARAGRAPH.CENTER, italic_default=True)
    new_para_before(anchor,
        "Hình 2.6. Sơ đồ ERD 2 — Đơn hàng & Thanh toán",
        alignment=WD_ALIGN_PARAGRAPH.CENTER, italic_default=True)
    new_para_before(anchor,
        "Sơ đồ ERD 2 tập trung vào luồng đặt hàng và thanh toán, gồm 4 thực "
        "thể chính: orders, order_items, vouchers, voucher_user_usages. "
        "Bảng orders đã được mở rộng với các trường người gửi (sender_name, "
        "sender_phone), người nhận (recipient_name, phone, shipping_address), "
        "ngày + khung giờ giao và lời nhắn in trên thiệp.",
        first_line_indent=Pt(18))

    # ===== Hình 2.7 — ERD 3 =====
    new_para_before(anchor, "[ Chèn ảnh sơ đồ ERD 3 ở đây ]",
                    alignment=WD_ALIGN_PARAGRAPH.CENTER, italic_default=True)
    new_para_before(anchor,
        "Hình 2.7. Sơ đồ ERD 3 — Nội dung & Giao tiếp",
        alignment=WD_ALIGN_PARAGRAPH.CENTER, italic_default=True)
    new_para_before(anchor,
        "Sơ đồ ERD 3 gồm 6 thực thể phục vụ quản trị nội dung blog, hệ thống "
        "chat 2 chiều khách – admin và cấu hình website: admins, "
        "contact_messages, contact_replies, blog_categories, blog_posts, "
        "website_settings. Đây là phần đứng độc lập tương đối với phần bán "
        "hàng — chỉ tham chiếu sang bảng users qua khoá ngoại trong "
        "contact_messages và contact_replies.",
        first_line_indent=Pt(18))

    # ===== Mối liên kết giữa 3 ERD =====
    new_para_before(anchor, [
        ("Mối liên kết giữa 3 sơ đồ ERD ", True),
        ("được thực hiện thông qua hai thực thể \"bridge\" là users và "
         "products — cả hai xuất hiện đầy đủ (với toàn bộ thuộc tính) "
         "trong ERD 1. Ở ERD 2 và ERD 3, hai thực thể này được vẽ lại "
         "dưới dạng hộp viền nét đứt với ghi chú [bridge ← ERD 1] để báo "
         "đó không phải bảng mới mà chỉ là tham chiếu sang sơ đồ khác. "
         "Cụ thể: ERD 2 dùng users.id (làm khoá ngoại trong orders, "
         "voucher_user_usages) và products.id (khoá ngoại trong "
         "order_items); ERD 3 dùng users.id (khoá ngoại trong "
         "contact_messages và contact_replies). Khi cần tra cứu chi "
         "tiết thuộc tính của users hoặc products, người đọc luôn quay "
         "về ERD 1 — nhờ vậy không có thông tin nào bị mất khi tách "
         "sơ đồ.", False),
    ], first_line_indent=Pt(18))

    print("  Đã thêm 3 caption Hình 2.5/2.6/2.7 + 4 đoạn giải thích vào section 2.3")


# ============================================================
# 4. Cập nhật Danh mục hình ảnh
# ============================================================
def update_danh_muc(doc):
    """- Renumber các entry Hình 2.X (X >= 5) → +3 (đã được làm bởi renumber_all)
       - Thêm 3 entry Hình 2.5/2.6/2.7 ERD vào đúng vị trí."""
    # Tìm vị trí chèn: sau Hình 2.4 trong Danh mục (paragraphs ~65)
    target = None
    for i, p in enumerate(doc.paragraphs):
        if i > 130: break
        if p.text.strip().startswith("Hình 2.4."):
            target = p; break
    if target is None:
        print("  Không tìm thấy Hình 2.4 trong Danh mục"); return

    # Idempotency check
    for p in doc.paragraphs[:130]:
        if "Hình 2.5. Sơ đồ ERD 1" in p.text:
            print("  Đã có Hình 2.5 ERD trong Danh mục — bỏ qua"); return

    entries = [
        "Hình 2.5. Sơ đồ ERD 1 — Khách hàng & Sản phẩm",
        "Hình 2.6. Sơ đồ ERD 2 — Đơn hàng & Thanh toán",
        "Hình 2.7. Sơ đồ ERD 3 — Nội dung & Giao tiếp",
    ]
    for entry in reversed(entries):
        new_p = OxmlElement('w:p')
        target._p.addnext(new_p)
        para = Paragraph(new_p, target._parent)
        run = para.add_run(entry)
        run.font.name = "Times New Roman"
        run.font.size = Pt(13)
        para.paragraph_format.line_spacing = 1.5
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    print(f"  Đã thêm 3 entry ERD (Hình 2.5/2.6/2.7) vào Danh mục hình ảnh")


def main():
    doc = Document(DOC)
    print("\n[1/4] Xoá heading 2.3.1 cũ:")
    remove_old_subsection(doc)
    print("\n[2/4] Renumber Hình 2.X (X≥5) → +3:")
    renumber_all(doc)
    print("\n[3/4] Thêm 3 caption ERD + giải thích vào section 2.3:")
    add_erd_captions(doc)
    print("\n[4/4] Cập nhật Danh mục hình ảnh:")
    update_danh_muc(doc)
    doc.save(DOC)
    print("\nOK — đã lưu lại file.")


if __name__ == "__main__":
    main()
