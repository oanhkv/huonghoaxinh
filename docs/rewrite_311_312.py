"""
Rewrite sections 3.1.1 (Mô hình MVC trong Laravel) and 3.1.2
(Áp dụng MVC vào dự án Hương Hoa Xinh) with a more student-friendly
tone, concrete project examples, and properly structured sub-blocks.

Style: Times New Roman 13pt, line spacing 1.5, justify body, bold
inline sub-labels.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

DOC = r"C:\Users\Kieu Anh\Desktop\CD1\Mã đề 18_Nhóm 3_Bài Thi.docx"

# --------------------------- helpers ---------------------------
def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def set_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), '808080')
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def new_para_before(anchor, segments, *, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                    line_spacing=1.5, first_line_indent=None,
                    space_before=None, space_after=None,
                    font_name="Times New Roman", size_pt=13):
    """segments: list of (text, bold?, italic?) tuples OR a plain string."""
    new_p = OxmlElement('w:p')
    anchor._p.addprevious(new_p)
    p = Paragraph(new_p, anchor._parent)
    if isinstance(segments, str):
        segments = [(segments, False, False)]
    for seg in segments:
        if len(seg) == 2:
            txt, bold = seg; italic = False
        else:
            txt, bold, italic = seg
        run = p.add_run(txt)
        run.font.name = font_name
        run.font.size = Pt(size_pt)
        run.bold = bold
        run.italic = italic
    p.alignment = alignment
    pf = p.paragraph_format
    pf.line_spacing = line_spacing
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    if space_before is not None:
        pf.space_before = space_before
    if space_after is not None:
        pf.space_after = space_after
    return p


def new_bullet_before(anchor, segments):
    """Bullet line — '•  ' prefix, justify, indent left."""
    if isinstance(segments, str):
        segments = [(segments, False, False)]
    segs = [("•  ", False, False)] + list(segments)
    p = new_para_before(anchor, segs,
                        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                        line_spacing=1.5)
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.first_line_indent = Cm(-0.6)
    return p


def new_code_block_before(anchor, code_text, doc):
    """1-cell table styled as a code block (Consolas, light grey fill)."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    anchor._p.addprevious(tbl._tbl)
    cell = tbl.rows[0].cells[0]
    cell.text = ""
    set_cell_shading(cell, "F5F5F5")
    set_cell_borders(cell)
    lines = code_text.rstrip("\n").split("\n")
    cell.paragraphs[0].text = lines[0]
    for line in lines[1:]:
        cell.add_paragraph(line)
    for cp in cell.paragraphs:
        cp.paragraph_format.line_spacing = 1.15
        cp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        cp.paragraph_format.space_before = Pt(0)
        cp.paragraph_format.space_after = Pt(0)
        cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for r in cp.runs:
            r.font.name = "Consolas"
            r.font.size = Pt(11)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    return tbl


# --------------------------- main ---------------------------
def main():
    doc = Document(DOC)

    # Locate 3.1.1, 3.1.2, 3.1.3 headings
    paras = list(doc.paragraphs)
    h311 = h312 = h313 = None
    for i, p in enumerate(paras):
        if p.style.name == "Heading 3":
            t = p.text.strip()
            if t.startswith("3.1.1"): h311 = (i, p)
            elif t.startswith("3.1.2"): h312 = (i, p)
            elif t.startswith("3.1.3"): h313 = (i, p)
    if None in (h311, h312, h313):
        raise SystemExit("Không tìm thấy đủ 3 heading 3.1.1 / 3.1.2 / 3.1.3")

    # Remove old body paragraphs of 3.1.1 (between h311 and h312)
    to_remove = []
    i311 = h311[0]; i312 = h312[0]; i313 = h313[0]
    for j in range(i311 + 1, i312):
        to_remove.append(paras[j])
    for j in range(i312 + 1, i313):
        to_remove.append(paras[j])
    for p in to_remove:
        p._p.getparent().remove(p._p)

    # ============================================================
    # 3.1.1 — new content (anchor = heading 3.1.2)
    # ============================================================
    anchor = h312[1]

    new_para_before(anchor,
        "Khi học môn Lập trình Web ở năm 3, nhóm em được làm quen với rất "
        "nhiều cách tổ chức mã nguồn khác nhau. Trong số đó, MVC – viết tắt "
        "của Model – View – Controller – là mẫu thiết kế phổ biến nhất và "
        "cũng chính là cách mà framework Laravel khuyến nghị sử dụng. Nói "
        "một cách dễ hiểu, MVC giúp chia một ứng dụng web thành ba phần "
        "riêng biệt, mỗi phần chỉ làm đúng việc của nó, từ đó tránh tình "
        "trạng \"code lẫn lộn\" rất khó bảo trì khi dự án lớn dần lên.",
        first_line_indent=Pt(18))

    new_para_before(anchor,
        "Cụ thể, ba thành phần của MVC trong Laravel hoạt động như sau:",
        first_line_indent=Pt(18))

    new_bullet_before(anchor, [
        ("Model", True), (" (mô hình dữ liệu) là lớp đại diện cho dữ liệu "
         "và các nghiệp vụ thao tác với dữ liệu. Trong Laravel, mỗi Model "
         "thường ứng với một bảng trong cơ sở dữ liệu MySQL và được xây "
         "dựng dựa trên Eloquent ORM. Ví dụ Model ", False),
        ("Product.php", False, True),
        (" ứng với bảng ", False),
        ("products", False, True),
        ("; muốn lấy danh sách hoa nổi bật chỉ cần gọi ", False),
        ("Product::where('is_featured', 1)->get()", False, True),
        (" thay vì viết câu SQL thủ công.", False),
    ])

    new_bullet_before(anchor, [
        ("View", True),
        (" (giao diện hiển thị) là phần mà người dùng nhìn thấy trên trình "
         "duyệt. Trong Laravel, View được viết bằng Blade Template "
         "(.blade.php) – sự kết hợp giữa HTML và cú pháp Blade gọn nhẹ "
         "như ", False),
        ("{{ $bien }}", False, True), (", ", False),
        ("@foreach", False, True),  (", ", False),
        ("@extends", False, True),  (", ", False),
        ("@section", False, True),
        (". View chỉ chịu trách nhiệm hiển thị, không xử lý logic phức "
         "tạp.", False),
    ])

    new_bullet_before(anchor, [
        ("Controller", True),
        (" (bộ điều phối) là cầu nối giữa Model và View. Khi người dùng "
         "gửi yêu cầu, Controller sẽ nhận yêu cầu đó, gọi Model để lấy "
         "dữ liệu, sau đó trả về View tương ứng kèm theo dữ liệu đã "
         "chuẩn bị sẵn.", False),
    ])

    new_para_before(anchor, [
        ("Để dễ hình dung, ", True),
        ("nhóm em thường so sánh MVC với một nhà hàng: khách hàng "
         "(trình duyệt) đặt món qua nhân viên phục vụ (Controller); "
         "nhân viên này không tự nấu mà chuyển yêu cầu xuống bếp "
         "(Model) – nơi nắm rõ công thức và nguyên liệu; sau khi món "
         "được nấu xong, nhân viên bưng món ra và trình bày đẹp mắt "
         "trên bàn ăn (View) cho khách thưởng thức.", False),
    ], first_line_indent=Pt(18))

    new_para_before(anchor, [
        ("Luồng xử lý một request trong Laravel ", True),
        ("diễn ra qua các bước: Trình duyệt gửi HTTP Request → Laravel "
         "tra cứu file ", False),
        ("routes/web.php", False, True),
        (" để biết phải gọi Controller nào → Controller gọi Model "
         "(Eloquent) để truy vấn cơ sở dữ liệu → Model trả dữ liệu về "
         "Controller → Controller chọn Blade View phù hợp và \"bind\" dữ "
         "liệu vào → Laravel render thành HTML rồi trả về cho trình "
         "duyệt.", False),
    ], first_line_indent=Pt(18))

    new_para_before(anchor,
        "Nhờ cách phân tách ba lớp rõ ràng này, việc sửa giao diện không "
        "lo ảnh hưởng đến logic nghiệp vụ; ngược lại, khi thay đổi cấu "
        "trúc cơ sở dữ liệu cũng không phải viết lại toàn bộ View. Đồng "
        "thời, mỗi thành viên trong nhóm có thể làm việc song song trên "
        "từng phần khác nhau (người làm giao diện, người làm Controller, "
        "người làm database) mà không bị chồng chéo code – đây là một "
        "lợi ích rất lớn khi làm việc nhóm trên Git.",
        first_line_indent=Pt(18))

    # ============================================================
    # 3.1.2 — new content (anchor = heading 3.1.3)
    # ============================================================
    anchor = h313[1]

    new_para_before(anchor,
        "Khi bắt tay vào xây dựng website Hương Hoa Xinh, nhóm em quyết "
        "định bám sát cấu trúc MVC mặc định của Laravel 11 và sắp xếp "
        "file theo đúng vị trí quy ước của framework. Cách làm này tuy "
        "ban đầu hơi gò bó nhưng càng về sau càng thấy lợi: khi cần "
        "thêm chức năng mới, mỗi người trong nhóm đều biết ngay nên đặt "
        "file ở đâu, đặt tên thế nào để cả nhóm cùng hiểu. Dưới đây là "
        "cách MVC được áp dụng cụ thể trong dự án:",
        first_line_indent=Pt(18))

    # ---- a) Model ----
    new_para_before(anchor, [
        ("a) Lớp Model – ", True),
        ("app/Models/", True, True),
    ], alignment=WD_ALIGN_PARAGRAPH.LEFT)

    new_para_before(anchor,
        "Toàn bộ 16 bảng trong cơ sở dữ liệu ban_hoa đều được ánh xạ "
        "sang 16 Model tương ứng. Nhờ Eloquent ORM, khi muốn truy vấn "
        "hay cập nhật dữ liệu, nhóm em chỉ cần gọi qua Model là đủ, "
        "không phải viết SQL thủ công. Một số Model tiêu biểu:",
        first_line_indent=Pt(18))

    bullets_model = [
        [("User.php", False, True), (", ", False),
         ("Admin.php", False, True),
         (" – tài khoản khách hàng và tài khoản quản trị (dùng 2 guard "
          "riêng biệt).", False)],
        [("Product.php", False, True), (", ", False),
         ("Category.php", False, True),
         (" – sản phẩm hoa và danh mục hoa.", False)],
        [("Order.php", False, True), (", ", False),
         ("OrderItem.php", False, True),
         (" – đơn hàng và chi tiết từng sản phẩm trong đơn.", False)],
        [("Cart.php", False, True), (", ", False),
         ("Wishlist.php", False, True),
         (" – giỏ hàng (hỗ trợ cả khách vãng lai qua session) và danh "
          "sách yêu thích.", False)],
        [("Voucher.php", False, True), (", ", False),
         ("VoucherUserUsage.php", False, True),
         (" – mã giảm giá và lượt sử dụng theo từng khách.", False)],
        [("BlogPost.php", False, True), (", ", False),
         ("BlogCategory.php", False, True),
         (" – bài viết blog và danh mục blog.", False)],
        [("ContactMessage.php", False, True), (", ", False),
         ("ContactReply.php", False, True), (", ", False),
         ("Review.php", False, True), (", ", False),
         ("WebsiteSetting.php", False, True),
         (" – tin nhắn liên hệ, phản hồi của admin, đánh giá sản phẩm "
          "và cài đặt website.", False)],
    ]
    for segs in bullets_model:
        new_bullet_before(anchor, segs)

    new_para_before(anchor,
        "Trong mỗi Model, nhóm em còn định nghĩa thêm các quan hệ "
        "(relationship) như hasMany, belongsTo. Ví dụ Model Product có "
        "hàm category() để truy xuất danh mục cha và reviews() để lấy "
        "danh sách đánh giá – nhờ đó chỉ với một dòng code là có thể "
        "lấy được sản phẩm kèm theo cả danh mục và toàn bộ đánh giá "
        "của nó.",
        first_line_indent=Pt(18))

    # ---- b) Controller ----
    new_para_before(anchor, [
        ("b) Lớp Controller – ", True),
        ("app/Http/Controllers/", True, True),
    ], alignment=WD_ALIGN_PARAGRAPH.LEFT)

    new_para_before(anchor,
        "Để dễ phân chia công việc giữa các thành viên và tránh việc "
        "trang khách hàng lẫn với trang admin, nhóm em chia thư mục "
        "Controllers thành hai nhánh riêng biệt:",
        first_line_indent=Pt(18))

    new_bullet_before(anchor, [
        ("Admin/", True, True),
        (" – chứa Controller cho khu vực quản trị: ", False),
        ("DashboardController", False, True), (", ", False),
        ("ProductController", False, True), (", ", False),
        ("CategoryController", False, True), (", ", False),
        ("OrderController", False, True), (", ", False),
        ("UserController", False, True), (", ", False),
        ("VoucherController", False, True), (", ", False),
        ("ReviewController", False, True), (", ", False),
        ("RevenueController", False, True), (", ", False),
        ("ContactMessageController", False, True), (", ", False),
        ("WebsiteSettingController", False, True), (", ", False),
        ("AdminProfileController", False, True),
        (". Mỗi Controller phụ trách một module CRUD riêng.", False),
    ])
    new_bullet_before(anchor, [
        ("Frontend/", True, True),
        (" – chứa Controller cho phía khách hàng: ", False),
        ("HomeController", False, True), (", ", False),
        ("ShopController", False, True), (", ", False),
        ("ProductController", False, True), (", ", False),
        ("CartController", False, True), (", ", False),
        ("PaymentController", False, True), (", ", False),
        ("WishlistController", False, True), (", ", False),
        ("BlogController", False, True), (", ", False),
        ("ProductReviewController", False, True), (", ", False),
        ("ShippingEstimateController", False, True),
        (".", False),
    ])
    new_bullet_before(anchor, [
        ("Auth/", True, True),
        (" – chứa Controller xác thực được sinh sẵn bởi Laravel Breeze "
         "(đăng ký, đăng nhập, quên mật khẩu, xác thực email).", False),
    ])

    # ---- c) View ----
    new_para_before(anchor, [
        ("c) Lớp View – ", True),
        ("resources/views/", True, True),
    ], alignment=WD_ALIGN_PARAGRAPH.LEFT)

    new_para_before(anchor,
        "Toàn bộ giao diện được viết bằng Blade Template và tổ chức "
        "thành 4 nhóm chính:",
        first_line_indent=Pt(18))

    new_bullet_before(anchor, [
        ("layouts/", True, True),
        (" – các template cha (", False),
        ("app.blade.php", False, True),
        (" cho frontend, ", False),
        ("admin.blade.php", False, True),
        (" cho admin) chứa header, footer, navigation. Các trang con "
         "chỉ cần ", False),
        ("@extends('layouts.app')", False, True),
        (" và ", False),
        ("@section('content')", False, True),
        (" là kế thừa được toàn bộ bố cục.", False),
    ])
    new_bullet_before(anchor, [
        ("frontend/", True, True),
        (" – giao diện khách hàng: trang chủ, shop, chi tiết sản phẩm, "
         "giỏ hàng, thanh toán, blog, wishlist, tài khoản, đặt hàng, "
         "viết đánh giá…", False),
    ])
    new_bullet_before(anchor, [
        ("admin/", True, True),
        (" – giao diện quản trị: dashboard, các trang CRUD của từng "
         "module, trang doanh thu, tin nhắn liên hệ, cài đặt website.", False),
    ])
    new_bullet_before(anchor, [
        ("auth/", True, True),
        (" – các form đăng nhập / đăng ký / quên mật khẩu / xác thực "
         "email.", False),
    ])

    # ---- d) Routes ----
    new_para_before(anchor, [
        ("d) Routes – ", True),
        ("routes/web.php", True, True),
    ], alignment=WD_ALIGN_PARAGRAPH.LEFT)

    new_para_before(anchor,
        "Đây là nơi nhóm em \"khai báo bản đồ\" của toàn bộ website. "
        "Mỗi URL được gắn với một method trong Controller cụ thể và đi "
        "kèm middleware bảo vệ tương ứng. Ví dụ:",
        first_line_indent=Pt(18))

    new_code_block_before(anchor,
        "Route::get('/shop', [ShopController::class, 'index'])->name('shop');\n"
        "\n"
        "Route::middleware('admin')->prefix('admin')->name('admin.')->group(function () {\n"
        "    Route::resource('products', ProductController::class);\n"
        "    Route::resource('orders',   OrderController::class);\n"
        "    Route::resource('vouchers', VoucherController::class);\n"
        "});",
        doc=doc)

    new_para_before(anchor,
        "Middleware admin đảm bảo chỉ tài khoản có guard admin mới truy "
        "cập được nhóm route /admin/*; tương tự middleware auth bảo vệ "
        "các route yêu cầu khách hàng đăng nhập như checkout, wishlist, "
        "đánh giá sản phẩm.",
        first_line_indent=Pt(18))

    # ---- e) Ví dụ minh hoạ luồng ----
    new_para_before(anchor,
        "e) Một ví dụ minh hoạ luồng MVC trong dự án",
        alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.5,
        first_line_indent=None,
    )
    # Bold the heading line we just added (manually set bold runs)
    # Re-style: locate it and bold runs
    # (alternative: pass [(text, True)] but we want left-align without indent)
    # Simpler: just set bold on its runs
    for p in doc.paragraphs:
        if p.text.strip() == "e) Một ví dụ minh hoạ luồng MVC trong dự án":
            for r in p.runs:
                r.bold = True
            break

    new_para_before(anchor,
        "Để thấy rõ MVC vận hành như thế nào trong thực tế, hãy theo "
        "dõi điều gì xảy ra khi khách hàng truy cập trang "
        "/shop?keyword=hồng:",
        first_line_indent=Pt(18))

    steps = [
        "1. Trình duyệt gửi yêu cầu GET tới máy chủ.",
        "2. Laravel đọc file routes/web.php, tìm thấy URL /shop được "
        "khai báo tương ứng với ShopController@index.",
        "3. ShopController@index gọi Model Product để truy vấn các sản "
        "phẩm có tên chứa \"hồng\" và đang được kích hoạt: "
        "Product::where('name','like','%hồng%')->where('is_active',1)"
        "->paginate(12).",
        "4. Eloquent (Model) trả về một collection các sản phẩm kèm "
        "theo thông tin danh mục (eager-loading qua with('category')).",
        "5. Controller bind dữ liệu này vào view frontend.shop và trả "
        "về.",
        "6. Blade render HTML cuối cùng và gửi về trình duyệt cho khách "
        "hàng nhìn thấy.",
    ]
    for s in steps:
        p = new_para_before(anchor, s,
                            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                            line_spacing=1.5)
        p.paragraph_format.left_indent = Cm(0.6)

    new_para_before(anchor,
        "Như vậy, mỗi thành phần trong MVC chỉ làm đúng việc của mình "
        "– Model lo dữ liệu, View lo hiển thị, Controller lo điều phối. "
        "Nhờ vậy, dự án Hương Hoa Xinh dễ bảo trì, dễ mở rộng và đặc "
        "biệt thuận tiện khi nhóm em phân chia công việc cho từng "
        "thành viên trên GitHub.",
        first_line_indent=Pt(18))

    doc.save(DOC)
    print("OK – đã viết lại mục 3.1.1 và 3.1.2")


if __name__ == "__main__":
    main()
