"""
Re-format the newly inserted 3.3 section + 2 tables + danh mục bảng biểu
entries to match the report style:
  - Font: Times New Roman, size 13
  - Line spacing: 1.5
  - Body paragraphs justified
  - Captions centered italic, table headers centered bold
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL

DOC_PATH = r"C:\Users\Kieu Anh\Desktop\CD1\Mã đề 18_Nhóm 3_Bài Thi.docx"


def set_para_format(p, *, alignment=None, line_spacing=1.5,
                    font_size=13, bold=False, italic=False,
                    color=None, font_name="Times New Roman"):
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.line_spacing = line_spacing
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    for r in p.runs:
        r.font.name = font_name
        r.font.size = Pt(font_size)
        if bold is not None:
            r.bold = bold
        if italic is not None:
            r.italic = italic
        if color is not None:
            r.font.color.rgb = color


def format_cell(cell, *, alignment, bold=False, italic=False,
                color=None, font_size=13):
    for p in cell.paragraphs:
        set_para_format(p, alignment=alignment, line_spacing=1.5,
                        font_size=font_size, bold=bold, italic=italic,
                        color=color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


GREEN = RGBColor(0x1B, 0x7A, 0x37)
RED   = RGBColor(0xB0, 0x21, 0x21)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def main():
    doc = Document(DOC_PATH)

    # --------- 1. Format 3.3 section paragraphs ----------
    target_idxs = {}
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t.startswith("3.3.") and "Kiểm thử" in t:
            target_idxs["heading"] = i
        elif t.startswith("Sau khi hoàn thiện các chức năng"):
            target_idxs["intro"] = i
        elif t == "Bảng 3.1. Tổng quan kết quả kiểm thử hệ thống":
            target_idxs["cap1"] = i
        elif t == "Bảng 3.2. Một số test case tiêu biểu và kết quả":
            target_idxs["cap2"] = i

    if "heading" in target_idxs:
        p = doc.paragraphs[target_idxs["heading"]]
        # keep Heading 2 style; just enforce line_spacing + font + size
        set_para_format(p, line_spacing=1.5, font_size=13, bold=True)

    if "intro" in target_idxs:
        p = doc.paragraphs[target_idxs["intro"]]
        set_para_format(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                        line_spacing=1.5, font_size=13)
        # Đầu dòng thụt vào theo phong cách báo cáo
        p.paragraph_format.first_line_indent = Pt(18)

    for key in ("cap1", "cap2"):
        if key in target_idxs:
            p = doc.paragraphs[target_idxs[key]]
            set_para_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                            line_spacing=1.5, font_size=13, italic=True)

    # --------- 2. Format 2 new tables (last two before phụ lục table) ----------
    # Tables before formatting: indexes 10 and 11 are the new ones (0-indexed)
    # Identify by header text to be safe
    new_tables = []
    for i, t in enumerate(doc.tables):
        if not t.rows: continue
        head = [c.text.strip() for c in t.rows[0].cells]
        if head[:2] == ["STT", "Module / Chức năng"]:
            new_tables.append(("summary", t))
        elif head[:2] == ["Mã TC", "Module"]:
            new_tables.append(("detail", t))

    for kind, t in new_tables:
        # Header row
        for c in t.rows[0].cells:
            format_cell(c, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                        bold=True, color=WHITE, font_size=13)
        # Data rows
        for ri, row in enumerate(t.rows[1:], start=1):
            cells = row.cells
            if kind == "summary":
                # STT / Module / Total / Pass / Fail / % Pass
                is_total = (ri == len(t.rows) - 1)
                format_cell(cells[0], alignment=WD_ALIGN_PARAGRAPH.CENTER,
                            bold=is_total, font_size=13)
                format_cell(cells[1],
                            alignment=(WD_ALIGN_PARAGRAPH.CENTER if is_total
                                       else WD_ALIGN_PARAGRAPH.LEFT),
                            bold=is_total, font_size=13)
                format_cell(cells[2], alignment=WD_ALIGN_PARAGRAPH.CENTER,
                            bold=is_total, font_size=13)
                format_cell(cells[3], alignment=WD_ALIGN_PARAGRAPH.CENTER,
                            bold=is_total, color=GREEN, font_size=13)
                format_cell(cells[4], alignment=WD_ALIGN_PARAGRAPH.CENTER,
                            bold=is_total, color=RED, font_size=13)
                format_cell(cells[5], alignment=WD_ALIGN_PARAGRAPH.CENTER,
                            bold=is_total, color=GREEN, font_size=13)
            else:
                # Mã / Module / Content / Expected / Actual / Status
                format_cell(cells[0], alignment=WD_ALIGN_PARAGRAPH.CENTER,
                            font_size=13)
                format_cell(cells[1], alignment=WD_ALIGN_PARAGRAPH.LEFT,
                            font_size=13)
                format_cell(cells[2], alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                            font_size=13)
                format_cell(cells[3], alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                            font_size=13)
                format_cell(cells[4], alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                            font_size=13)
                # Status cell — color theo Pass/Fail
                status_text = cells[5].text.strip()
                clr = GREEN if status_text.lower() == "pass" else RED
                format_cell(cells[5], alignment=WD_ALIGN_PARAGRAPH.CENTER,
                            bold=True, color=clr, font_size=13)

    # --------- 3. Format 'Danh mục bảng biểu' new entries ----------
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if (t.startswith("Bảng 3.1") or t.startswith("Bảng 3.2")) and i < 200:
            # entries in TOC of tables (early in doc, before chương 1)
            set_para_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                            line_spacing=1.5, font_size=13)

    doc.save(DOC_PATH)
    print("OK")


if __name__ == "__main__":
    main()
